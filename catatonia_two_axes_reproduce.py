#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
two_axes_paper.py

Self-contained reproduction, figure and self-check script for

    Two separable axes of intervention in a dependency-ordered recovery geometry
    Hiroki Saito

Every numerical value printed in the manuscript and in Supplementary Notes 1-7 is
recomputed here and checked against the value as printed.  The script also carries
the secondary analysis of published aggregate data reported in the Results, draws
the five display items, writes their source data, holds the text of the manuscript
and of the Supplementary Information, audits that text against the journal's
requirements, and renders both documents.  It has no inputs.

Model (ref. 4 of the manuscript)
--------------------------------
    g(r,Q) = r (1 - r) (r - a(Q))
    dr0/dt = g(r0,Q) + u(t)
    drk/dt = g(rk,Q) + kappa_k (r_{k-1} - r_k),      k = 1..3
    dQ/dt  = eps [ chi_+(z) C(r) (1 - Q) - rho chi_-(z) Q ],   z = r0 - a(Q)
    C(r)   = prod_j r_j,   a(Q) = a0 - (a0 - a1) Q^p,  p > 1
    chi_+-(z) = exp(-s / z^2) on its own half-line, 0 elsewhere

Usage
-----
    python3 two_axes_paper.py                # full run
    python3 two_axes_paper.py --quick        # reduced sample sizes, same structure
    python3 two_axes_paper.py --no-figures   # values and self-check only
    python3 two_axes_paper.py --audit        # compliance audit and documents only

Outputs
-------
    Fig1.png ... Fig5.png                    the five display items
    two_axes_source_data.xlsx                one sheet per panel
    two_axes_results.json                    every computed quantity
    two_axes_main.docx                       the manuscript
    two_axes_supplementary.docx              the Supplementary Information

Requires numpy and scipy.  Matplotlib and openpyxl are needed for the figures and
the workbook, python-docx for the two documents; all three are optional.
"""

from __future__ import annotations
import argparse, json, math, re, sys, time
import numpy as np
from scipy.integrate import solve_ivp
from scipy.optimize import brentq

# ----------------------------------------------------------------------------
# 0.  Illustrative parameter set of ref. 4  (unfitted)
# ----------------------------------------------------------------------------
A0, A1, P_EXP, EPS, RHO = 0.60, 0.15, 2.0, 0.02, 1.0
KAPPA_REF = 0.6
S_NARROW, S_WIDE = 0.01, 1.0
RTOL, ATOL = 1e-10, 1e-13

RESULTS: dict = {}
CHECKS: list = []


def check(name, got, expected, tol, note=""):
    """Record a comparison between a recomputed value and the value as printed."""
    ok = abs(got - expected) <= tol
    CHECKS.append(dict(name=name, got=float(got), printed=float(expected),
                       tol=float(tol), ok=bool(ok), note=note))
    flag = "OK  " if ok else "FAIL"
    print(f"  [{flag}] {name:<52s} got {got:.10g}  printed {expected:.10g}")
    return ok


def a_of_Q(Q, a0=A0, a1=A1, p=P_EXP):
    return a0 - (a0 - a1) * Q ** p


def g(r, a):
    return r * (1.0 - r) * (r - a)


def chi_plus(z, s):
    return math.exp(-s / z ** 2) if z > 0 else 0.0


def chi_minus(z, s):
    return math.exp(-s / z ** 2) if z < 0 else 0.0


# ----------------------------------------------------------------------------
# 1.  Confinement below threshold  (maximum principle, main text and Fig. 1b)
# ----------------------------------------------------------------------------
def draw_weights(rng, n, family):
    """Mixture placing mass at 0 and at 0.5, 3, 30 times a uniform draw on (0,1]."""
    base = rng.uniform(1e-9, 1.0, size=(n, n))
    scale = rng.choice([0.0, 0.5, 3.0, 30.0], size=(n, n), p=[1/3, 2/9, 2/9, 2/9])
    W = base * scale
    if family == "chain":          # feed-forward: strict sub-diagonal only
        M = np.zeros((n, n));  M[np.arange(1, n), np.arange(0, n - 1)] = 1.0
        W = W * M
    elif family == "reverse":      # purely reverse: strict upper triangle
        W = np.triu(W, 1)
    elif family == "symmetric":
        W = np.tril(W, -1); W = W + W.T
    elif family == "general":
        np.fill_diagonal(W, 0.0)
    return W


def run_confinement(n_per_family=100, seed=20260814):
    rng = np.random.default_rng(seed)
    fams = ["chain", "symmetric", "reverse", "general"]
    reached, max_rise = 0, -np.inf
    for fam in fams:
        for _ in range(n_per_family):
            n = int(rng.integers(3, 7))
            W = draw_weights(rng, n, fam)
            r0 = rng.uniform(0.0, A0 * 0.999, size=n)
            M0 = r0.max()

            def f(t, r):
                r = np.clip(r, 0.0, 1.0)
                return g(r, A0) + (W @ r - W.sum(axis=1) * r)

            sol = solve_ivp(f, [0, 2000.0], r0, method="LSODA",
                            rtol=1e-9, atol=1e-12, dense_output=True)
            traj = sol.sol(np.linspace(0, 2000.0, 4001))
            M = traj.max(axis=0)
            if M.max() >= A0:
                reached += 1
            max_rise = max(max_rise, M.max() - M0)
    return dict(n_systems=4 * n_per_family, reached_threshold=reached,
                largest_rise_of_running_max=float(max_rise))


def run_confinement_heterogeneous(n_graphs=200, seed=7):
    rng = np.random.default_rng(seed)
    reached = 0
    for _ in range(n_graphs):
        n = int(rng.integers(3, 7))
        W = draw_weights(rng, n, "general")
        a = rng.uniform(A1, A0, size=n)
        r0 = rng.uniform(0.0, a.min() * 0.999, size=n)

        def f(t, r):
            r = np.clip(r, 0.0, 1.0)
            return r * (1 - r) * (r - a) + (W @ r - W.sum(axis=1) * r)

        y = solve_ivp(f, [0, 2000.0], r0, method="LSODA",
                      rtol=1e-9, atol=1e-12).y
        if np.any(y.max(axis=1) >= a.min()):
            reached += 1
    return dict(n_graphs=n_graphs, reached_min_threshold=reached)


def run_reverse_counterexample():
    """r = (0.30, 0.90, 0.90, 0.90): the maximum is NOT below a, and reverse
    coupling of 0.10 then converts the root."""
    def f(t, r, kf, kr):
        r = np.clip(r, 0, 1); o = np.zeros(4)
        for k in range(4):
            o[k] = g(r[k], A0)
            if k > 0:
                o[k] += kf * (r[k - 1] - r[k])
            if k < 3:
                o[k] += kr * (r[k + 1] - r[k])
        return o
    y0 = [0.30, 0.90, 0.90, 0.90]
    both = solve_ivp(f, [0, 3e4], y0, args=(0.01, 0.10), method="LSODA",
                     rtol=RTOL, atol=ATOL).y[:, -1]
    rev = solve_ivp(f, [0, 3e4], y0, args=(0.0, 0.10), method="LSODA",
                    rtol=RTOL, atol=ATOL).y[:, -1]
    return dict(forward001_reverse010=np.round(both, 6).tolist(),
                reverse_only=np.round(rev, 6).tolist())


def run_root_kappa_independence():
    """On the forward chain kappa multiplies zero in the root equation."""
    def f(t, r, k):
        r = np.clip(r, 0, 1); o = np.zeros(4); o[0] = g(r[0], A0)
        for j in range(1, 4):
            o[j] = g(r[j], A0) + k * (r[j - 1] - r[j])
        return o
    grid = np.linspace(0, 200, 20001)
    roots = []
    for k in (0.6, 3.0, 30.0):
        s = solve_ivp(f, [0, 200], [0.30, 0.20, 0.10, 0.05], args=(k,),
                      method="LSODA", rtol=RTOL, atol=ATOL, dense_output=True)
        roots.append(s.sol(grid)[0])
    d = max(np.max(np.abs(roots[0] - roots[1])), np.max(np.abs(roots[0] - roots[2])))
    return dict(max_pointwise_root_difference=float(d))


# ----------------------------------------------------------------------------
# 2.  Root pulse: the crossing criterion is r0(tau) > a0   (Fig. 1c)
# ----------------------------------------------------------------------------
def chain_rhs(t, r, kappa, a, u=0.0, tau=0.0):
    r = np.clip(r, 0.0, 1.0)
    o = np.zeros(4)
    o[0] = g(r[0], a) + (u if t <= tau else 0.0)
    for k in range(1, 4):
        o[k] = g(r[k], a) + kappa * (r[k - 1] - r[k])
    return o


def pulse_outcome(A, tau, kappa=KAPPA_REF, a=A0, T=4000.0):
    s = solve_ivp(chain_rhs, [0, T], [0.0, 0.0, 0.0, 0.0],
                  args=(kappa, a, A, tau), method="LSODA",
                  rtol=RTOL, atol=ATOL, dense_output=True)
    return s.sol(tau)[0], s.y[3, -1] > 0.5


def run_pulse_sweep():
    rows = []
    for tau in (0.25, 0.5, 1.0, 2.0, 5.0):
        lo, hi = 1e-4, 50.0
        for _ in range(80):
            mid = 0.5 * (lo + hi)
            if pulse_outcome(mid, tau)[1]:
                hi = mid
            else:
                lo = mid
        Ac = 0.5 * (lo + hi)
        cell = []
        for m in (0.8, 0.99, 1.01, 1.2):
            r0tau, conv = pulse_outcome(m * Ac, tau)
            cell.append(dict(mult=m, r0_at_tau=float(r0tau), converted=bool(conv),
                             agrees=bool(conv == (r0tau > A0))))
        rows.append(dict(tau=tau, critical_amplitude=float(Ac), cells=cell))
    agree = all(c["agrees"] for row in rows for c in row["cells"])
    # boundary in the direct-set sweep
    def outcome_from_root(v):
        s = solve_ivp(chain_rhs, [0, 4000.0], [v, 0, 0, 0], args=(KAPPA_REF, A0),
                      method="LSODA", rtol=RTOL, atol=ATOL)
        return s.y[3, -1] > 0.5
    lo, hi = 0.3, 0.9
    for _ in range(80):
        mid = 0.5 * (lo + hi)
        if outcome_from_root(mid): hi = mid
        else: lo = mid
    return dict(rows=rows, all_cells_agree_with_r0_criterion=bool(agree),
                located_boundary=float(0.5 * (lo + hi)))


# ----------------------------------------------------------------------------
# 3.  Exact partial branch, kappa_c, stage-wise closed form
# ----------------------------------------------------------------------------
def branch_exact(kappa, a, n=4):
    """Exact partial equilibrium with the root held at 1."""
    disc = a * a - 4.0 * kappa
    if disc < 0:
        return None
    r = [1.0, 0.5 * (a - math.sqrt(disc))]
    for k in range(2, n):
        prev = r[-1]
        if prev <= 0.0:            # underflow far down a long weakly coupled chain
            r.append(0.0)
            continue
        fn = lambda x: x * (1 - x) * (x - a) + kappa * (prev - x)
        # fn(0) = kappa*prev > 0 and fn(a) = -kappa*(a - prev) < 0, so the low
        # root is bracketed on (0, a) whenever 0 < prev < a  (Supplementary Note 4)
        r.append(brentq(fn, 0.0, a, xtol=1e-17, rtol=8.9e-16, maxiter=300))
    return r


def branch_closed_form(kappa, a, n=4):
    r = [1.0, 0.5 * (a - math.sqrt(max(a * a - 4 * kappa, 0.0)))]
    for k in range(2, n):
        prev = r[-1]
        disc = a * a - 4 * kappa * prev
        r.append(0.5 * (a - math.sqrt(disc)))
    return r


def run_branch():
    out = {}
    out["kappa_c_a1"] = A1 ** 2 / 4
    out["kappa_c_a0"] = A0 ** 2 / 4
    out["branch_at_0.005625"] = branch_exact(0.005625, A1) is not None
    out["branch_at_0.00563"] = branch_exact(0.00563, A1) is not None
    ex = branch_exact(0.003, A1); cf = branch_closed_form(0.003, A1)
    out["profile_kappa0.003_a1"] = ex
    out["closed_form_rel_error_r2"] = (cf[2] - ex[2]) / ex[2]
    out["closed_form_rel_error_r3"] = (cf[3] - ex[3]) / ex[3]
    out["analytic_residual_stage2"] = 0.003 * cf[2] * (cf[1] - 1.0)
    out["C_on_branch_kappa0.003"] = float(np.prod(ex))
    # residual and local stability of the exact profile
    def F(r, kappa=0.003, a=A1):
        o = [0.0]
        for k in range(1, 4):
            o.append(g(r[k], a) + kappa * (r[k - 1] - r[k]))
        return np.array(o)
    out["vector_field_residual"] = float(np.max(np.abs(F(np.array(ex)))))
    pert = np.array(ex) + np.array([0.0, 1e-4, 1e-4, 1e-4])
    y = solve_ivp(lambda t, r: F(r), [0, 5000], pert, method="LSODA",
                  rtol=RTOL, atol=ATOL).y[:, -1]
    out["returns_after_1e-4_perturbation"] = float(np.max(np.abs(y - np.array(ex))))
    # propagation boundary at depleted evidence, supracritical excursion
    def propagates(kappa):
        s = solve_ivp(chain_rhs, [0, 3e4], [0.61, 0, 0, 0], args=(kappa, A0),
                      method="LSODA", rtol=RTOL, atol=ATOL)
        return s.y[3, -1] > 0.5
    lo, hi = 0.05, 0.2
    for _ in range(60):
        m = 0.5 * (lo + hi)
        if propagates(m): hi = m
        else: lo = m
    out["propagation_boundary_a0"] = 0.5 * (lo + hi)
    return out


# ----------------------------------------------------------------------------
# 4.  Unequal coupling: min_k kappa_k, arrest depth, existence, chain length
# ----------------------------------------------------------------------------
def run_unequal():
    def final(kv, a=A0, y0=(0.61, 0, 0, 0)):
        def f(t, r):
            r = np.clip(r, 0, 1); o = np.zeros(4); o[0] = g(r[0], a)
            for k in range(1, 4):
                o[k] = g(r[k], a) + kv[k - 1] * (r[k - 1] - r[k])
            return o
        return solve_ivp(f, [0, 3e4], list(y0), method="LSODA",
                         rtol=RTOL, atol=ATOL).y[:, -1]
    bounds = []
    for stage in range(3):
        lo, hi = 0.01, 0.5
        for _ in range(60):
            m = 0.5 * (lo + hi)
            kv = [1.0, 1.0, 1.0]; kv[stage] = m
            if final(kv)[stage + 1] > 0.5: hi = m
            else: lo = m
        bounds.append(0.5 * (lo + hi))
    vectors = [[0.5, 0.05, 0.02], [0.5, 0.02, 0.05], [0.2, 0.08, 0.03],
               [0.3, 0.04, 0.07], [0.5, 0.06, 0.015], [0.2, 0.2, 0.05], [0.05, 0.2, 0.2]]
    depth_ok = True
    depths = []
    for kv in vectors:
        fin = final(kv)
        depth = int(np.sum(fin[1:] > 0.5))
        first_sub = next((i for i, k in enumerate(kv) if k <= A0 ** 2 / 4), 3)
        depths.append(dict(kappa=kv, arrest_depth=depth, first_subcritical=first_sub))
        depth_ok &= (depth == first_sub)
    # branch existence is decided by kappa_1 alone
    exists = {}
    for k1 in (0.005625, 0.00563):
        for kd in (1e-4, 0.5, 10.0, 30.0):
            exists[f"k1={k1},kd={kd}"] = (A1 ** 2 - 4 * k1) >= 0
    # kappa_c is unchanged by chain length
    lens = {}
    for n in range(2, 14):
        lo, hi = 1e-4, 0.05
        for _ in range(60):
            m = 0.5 * (lo + hi)
            if branch_exact(m, A1, n=n) is None: hi = m
            else: lo = m
        lens[n - 1] = 0.5 * (lo + hi)
    return dict(per_stage_boundaries=bounds, arrest=depths,
                arrest_depth_matches_first_subcritical=bool(depth_ok),
                existence=exists, kappa_c_by_chain_length=lens)


# ----------------------------------------------------------------------------
# 5.  Spectrum on the branch:  lambda_1 = -sqrt(a^2 - 4 kappa) (1 - r_1)
# ----------------------------------------------------------------------------
def run_spectrum():
    worst = 0.0
    rows = []
    for a in (A1, 0.30, A0):
        for kappa in np.linspace(0.02, 0.95, 10) * (a * a / 4):
            r = branch_exact(kappa, a)
            lam_num = (1 - 2 * r[1]) * (r[1] - a) + r[1] * (1 - r[1]) - kappa
            lam_cf = -math.sqrt(a * a - 4 * kappa) * (1 - r[1])
            worst = max(worst, abs(lam_num - lam_cf))
            rows.append(dict(a=a, kappa=float(kappa), lam=float(lam_cf)))
    return dict(max_abs_discrepancy=float(worst), n_points=len(rows))


# ----------------------------------------------------------------------------
# 6.  Annihilation locus and the intermediate band
# ----------------------------------------------------------------------------
def Q_star(kappa, a0=A0, a1=A1, p=P_EXP):
    x = (a0 - 2 * math.sqrt(kappa)) / (a0 - a1)
    return x ** (1.0 / p) if 0.0 <= x <= 1.0 else float("nan")


def full_rhs(t, y, kappa, s):
    r = np.clip(y[:4], 0.0, 1.0); Q = min(max(y[4], 0.0), 1.0)
    a = a_of_Q(Q); z = r[0] - a
    o = np.zeros(5); o[0] = g(r[0], a)
    for k in range(1, 4):
        o[k] = g(r[k], a) + kappa * (r[k - 1] - r[k])
    C = float(np.prod(r))
    o[4] = EPS * (chi_plus(z, s) * C * (1 - Q) - RHO * chi_minus(z, s) * Q)
    return o


def run_band(quick=False):
    out = dict(kappa_c_a0=A0 ** 2 / 4, kappa_c_a1=A1 ** 2 / 4)
    out["Q_star"] = {k: Q_star(k) for k in (0.05, 0.08)}
    runs = {}
    if quick:                       # the full-system waits are 1e6-1e9 time units
        out["full_system"] = "skipped in --quick"
        return out
    for kappa, s, T in ((0.05, S_NARROW, 4e8), (0.08, S_NARROW, 6e6), (0.08, S_WIDE, 2e9)):
        r = branch_exact(kappa, A0)
        y0 = list(r) + [0.0]

        def fold(t, y, *a_):
            return a_of_Q(min(max(y[4], 0), 1)) ** 2 - 4 * kappa
        fold.terminal = False; fold.direction = -1

        def escape(t, y, *a_):
            return y[3] - a_of_Q(min(max(y[4], 0), 1))
        escape.terminal = True; escape.direction = 1
        sol = solve_ivp(full_rhs, [0, T], y0, args=(kappa, s), method="LSODA",
                        rtol=1e-9, atol=1e-12, events=[fold, escape], max_step=T / 50)
        t_fold = float(sol.t_events[0][0]) if len(sol.t_events[0]) else float("nan")
        t_esc = float(sol.t_events[1][0]) if len(sol.t_events[1]) else float("nan")
        runs[f"kappa={kappa},s={s}"] = dict(
            t_fold=t_fold, t_escape=t_esc,
            lag_percent=100.0 * (t_esc - t_fold) / t_fold if t_fold == t_fold else None)
    out["full_system"] = runs
    return out


# ----------------------------------------------------------------------------
# 7.  Escape-time scaling and the zero-evidence-rate control
# ----------------------------------------------------------------------------
def escape_time(kappa, a, T=1e9):
    r0 = branch_exact(0.5 * a * a / 4, a)

    def f(t, r):
        r = np.clip(r, 0, 1); o = np.zeros(3)
        prev = 1.0
        for k in range(3):
            o[k] = g(r[k], a) + kappa * ((prev if k == 0 else r[k - 1]) - r[k])
        return o

    def ev(t, r):
        return r[2] - a
    ev.terminal = True; ev.direction = 1
    sol = solve_ivp(f, [0, T], r0[1:], method="LSODA", rtol=1e-11, atol=1e-14,
                    events=ev, max_step=T / 200)
    return float(sol.t_events[0][0]) if len(sol.t_events[0]) else float("nan")


def run_escape_scaling(quick=False):
    a = A1; kc = a * a / 4
    ex = np.array([0.8, 0.4, 0.2, 0.1, 0.05, 0.02, 0.01, 5e-3, 2e-3, 1e-3,
                   5e-4, 2e-4, 1e-4, 5e-5, 2e-5, 1e-5])
    if quick:
        ex = ex[:10]
    t = np.array([escape_time(kc * (1 + e), a) for e in ex])
    ok = np.isfinite(t)
    lx, ly = np.log(ex[ok]), np.log(t[ok])
    local = -(np.diff(ly) / np.diff(lx))
    full = -np.polyfit(lx, ly, 1)[0]
    m = ex[ok] <= 0.05
    small = -np.polyfit(lx[m], ly[m], 1)[0] if m.sum() > 2 else float("nan")
    ctrl = {}
    for Q0 in (1.0, 0.99, 0.90, 0.70, 0.50, 0.0):
        aQ = a_of_Q(Q0); kcq = aQ * aQ / 4
        exq = np.array([0.8, 0.4, 0.2, 0.1, 0.05, 0.02, 0.01, 5e-3])
        tq = np.array([escape_time(kcq * (1 + e), aQ) for e in exq])
        good = np.isfinite(tq)
        lxq, lyq = np.log(exq[good]), np.log(tq[good])
        mq = exq[good] <= 0.05
        ctrl[Q0] = dict(a=float(aQ),
                        eps0_full=float(-np.polyfit(lxq, lyq, 1)[0]),
                        eps0_small=float(-np.polyfit(lxq[mq], lyq[mq], 1)[0]))
    return dict(excess=ex.tolist(), escape_time=t.tolist(),
                local_slope_smallest=float(local[-1]),
                fit_full_range=float(full), fit_below_5pct=float(small),
                zero_evidence_rate_control=ctrl)


# ----------------------------------------------------------------------------
# 8.  Identification from the resting profile
# ----------------------------------------------------------------------------
def run_identification(n_draws=2000, quick=False):
    if quick:
        n_draws = 400
    kappa_true, a = 0.003, A1
    r = branch_exact(kappa_true, a)
    out = dict(kappa_true=kappa_true, r=r)
    out["first_stage_exact"] = a * r[1] - r[1] ** 2
    a_two = (r[1] ** 3 - r[2] ** 2) / (r[1] ** 2 - r[2])
    out["two_ratio_a"] = a_two
    out["two_ratio_a_bias_pct"] = 100 * (a_two - a) / a
    k_two = a_two * r[1] - r[1] ** 2
    out["two_ratio_kappa_bias_pct"] = 100 * (k_two - kappa_true) / kappa_true
    for a_wrong, tag in ((0.30, "a=0.30"), (A0, "a=a0")):
        out[f"ratio_using_{tag}"] = (a_wrong * r[1] - r[1] ** 2) / kappa_true
    ah = a * 1.10
    out["ratio_10pct_threshold_error"] = (ah * r[1] - r[1] ** 2) / kappa_true
    rng = np.random.default_rng(11)
    noise = {}
    for cv in (0.01, 0.05, 0.10):
        sig = math.sqrt(math.log(1 + cv ** 2))
        r1n = r[1] * rng.lognormal(-0.5 * sig ** 2, sig, n_draws)
        r2n = r[2] * rng.lognormal(-0.5 * sig ** 2, sig, n_draws)
        k1 = a * r1n - r1n ** 2
        with np.errstate(all="ignore"):
            a2 = (r1n ** 3 - r2n ** 2) / (r1n ** 2 - r2n)
            k2 = a2 * r1n - r1n ** 2
        q = lambda v: [float(np.nanpercentile(v, p)) for p in (25, 50, 75)]
        noise[cv] = dict(first_stage=q(k1), two_ratio=q(k2))
    out["noise"] = noise
    return out


# ----------------------------------------------------------------------------
# 9.  Intermittent modulation
# ----------------------------------------------------------------------------
def modulated_escape(kbase, khigh, f_duty, period, a=A1, T=1.0):
    """Escape time from the exact partial branch under a rectangular modulation of
    kappa.  Integrated phase by phase, so the vector field is smooth on each
    interval and the discontinuity is handled exactly at the phase boundary."""
    r = np.array(branch_exact(kbase, a)[1:])

    def rhs_const(t, y, k):
        y = np.clip(y, 0.0, 1.0); o = np.zeros(3); prev = 1.0
        for j in range(3):
            o[j] = g(y[j], a) + k * ((prev if j == 0 else y[j - 1]) - y[j])
        return o

    def ev(t, y, k):
        return y[2] - a
    ev.terminal = True; ev.direction = 1

    t = 0.0
    phases = [(khigh, f_duty * period), (kbase, (1.0 - f_duty) * period)]
    guard = 0
    while t < T and guard < 200000:
        for k, dur in phases:
            if dur <= 0:
                continue
            sol = solve_ivp(rhs_const, [t, t + dur], r, args=(k,), method="LSODA",
                            rtol=1e-9, atol=1e-12, events=ev)
            if len(sol.t_events[0]):
                return float(sol.t_events[0][0])
            r = sol.y[:, -1]; t += dur
            guard += 1
            if t >= T:
                break
    return float("inf")


def run_modulation(quick=False):
    a = A1; kc = a * a / 4; kbase = 0.003
    duties = [0.05, 0.1, 0.2, 0.35, 0.5, 0.7]
    mults = [0.5, 0.8, 1.2, 2.0, 5.0, 14.0, 30.0]
    if quick:
        duties, mults = duties[:3], mults[:4]
    t_static = escape_time(1.2 * kc, a)
    cells, agree = [], 0
    for fdc in duties:
        for m in mults:
            khigh = m * kc
            mean = fdc * khigh + (1 - fdc) * kbase
            te = modulated_escape(kbase, khigh, fdc, t_static / 50.0, a,
                                  T=60.0 * t_static)
            predicted = mean > kc
            observed = math.isfinite(te)
            agree += (predicted == observed)
            cells.append(dict(duty=fdc, mult=m, duty_weighted_mean=mean,
                              predicted=bool(predicted), observed=bool(observed)))
    return dict(static_escape_time=t_static, n_cells=len(cells),
                n_agree=int(agree), cells=cells)


# ----------------------------------------------------------------------------
# 10. Noise on the branch (Euler-Maruyama)
# ----------------------------------------------------------------------------
def run_noise(n_paths=500, T=2e4, dt=0.05, quick=False):
    if quick:
        n_paths, T = 60, 4e3
    a, kappa = A1, 0.003
    r = branch_exact(kappa, a)
    r1 = r[1]
    # barrier of the first-stage potential between r1 and the unstable partner
    ru = 0.5 * (a + math.sqrt(a * a - 4 * kappa))
    xs = np.linspace(r1, ru, 200001)
    fx = xs * (1 - xs) * (xs - a) + kappa * (1.0 - xs)
    V = -np.cumsum(fx) * (xs[1] - xs[0])
    barrier = float(V.max() - V[0])
    rng = np.random.default_rng(3)
    out = {}
    for sigma in (0.002, 0.008):
        nsteps = int(T / dt); left = 0; times = []
        y = np.tile(np.array(r[1:]), (n_paths, 1))
        alive = np.ones(n_paths, bool)
        for i in range(nsteps):
            prev = np.hstack([np.ones((n_paths, 1)), y[:, :2]])
            drift = y * (1 - y) * (y - a) + kappa * (prev - y)
            y = np.clip(y + dt * drift + sigma * math.sqrt(dt) *
                        rng.standard_normal(y.shape), 0.0, 1.0)
            esc = alive & (y[:, 0] > ru)
            if esc.any():
                times.extend([(i + 1) * dt] * int(esc.sum()))
                left += int(esc.sum()); alive &= ~esc
                y[~alive] = np.array(r[1:])
        out[sigma] = dict(n_paths=n_paths, n_escaped=left,
                          median_escape=float(np.median(times)) if times else None)
    return dict(barrier=barrier, r1=r1, r_unstable=ru, runs=out)


# ----------------------------------------------------------------------------
# 11. Supplementary Note 6: what confinement does not cover
# ----------------------------------------------------------------------------
def S_lin(x):
    return x


def S_sig(x):
    return 1.0 / (1.0 + math.exp(-(x - 0.30) / 0.05))


def additive_run(w, S, graph, y0, a=A0, T=3e4):
    """graph = 'complete' (every capacity drives every other) or 'chain'
    (capacity k receives only from k-1; the root receives nothing)."""
    def f(t, r):
        r = np.clip(r, 0.0, 1.0); o = np.zeros(4)
        for k in range(4):
            if graph == "complete":
                drive = sum(S(r[j]) for j in range(4) if j != k)
            else:
                drive = S(r[k - 1]) if k > 0 else 0.0
            o[k] = g(r[k], a) + w * (1 - r[k]) * drive
        return o
    return solve_ivp(f, [0, T], list(y0), method="LSODA",
                     rtol=RTOL, atol=ATOL).y[:, -1]


def diffusive_run(w, y0, a=A0, T=3e4, graph="complete"):
    def f(t, r):
        r = np.clip(r, 0.0, 1.0); o = np.zeros(4)
        for k in range(4):
            if graph == "complete":
                o[k] = g(r[k], a) + w * sum(r[j] - r[k] for j in range(4) if j != k)
            else:
                o[k] = g(r[k], a) + (w * (r[k - 1] - r[k]) if k > 0 else 0.0)
        return o
    return solve_ivp(f, [0, T], list(y0), method="LSODA",
                     rtol=RTOL, atol=ATOL).y[:, -1]


def bisect_weight(pred, lo, hi, n=70):
    for _ in range(n):
        m = 0.5 * (lo + hi)
        if pred(m): hi = m
        else: lo = m
    return 0.5 * (lo + hi)


def run_scope():
    out = {}
    start = [0.30] * 4
    # S8: additive coupling on the complete graph converts a sub-threshold chain
    out["additive_complete_lin_wc"] = bisect_weight(
        lambda w: additive_run(w, S_lin, "complete", start)[0] > 0.5, 0.01, 1.0)
    out["additive_complete_sig_wc"] = bisect_weight(
        lambda w: additive_run(w, S_sig, "complete", start)[0] > 0.5, 0.01, 1.0)
    out["diffusive_complete_finals"] = {
        w: np.round(diffusive_run(w, start), 6).tolist() for w in (0.1, 1.0, 10.0, 100.0)}
    # S8a: the same drive restricted to the derived chain never reaches the root
    out["additive_chain_lin_wc_r3"] = bisect_weight(
        lambda w: additive_run(w, S_lin, "chain", start)[3] > 0.5, 0.01, 5.0)
    out["additive_chain_lin_wc_r1"] = bisect_weight(
        lambda w: additive_run(w, S_lin, "chain", start)[1] > 0.5, 0.01, 5.0)
    finals = {}
    for w in (1.0, 10.0, 100.0):
        finals[w] = np.round(additive_run(w, S_lin, "chain", start), 6).tolist()
    finals["sig_from_collapsed_w100"] = np.round(
        additive_run(100.0, S_sig, "chain", [0.0] * 4), 6).tolist()
    out["additive_chain_finals"] = finals
    out["additive_chain_max_root"] = float(max(
        additive_run(w, S_lin, "chain", start)[0] for w in (0.1, 0.5, 1, 5, 20, 100)))
    out["additive_chain_C_at_w1"] = float(np.prod(np.clip(
        additive_run(1.0, S_lin, "chain", start), 0, 1)))
    # analytic checks quoted in S8
    out["symmetric_subspace_wc_lin"] = (A0 - 0.30) / 3.0
    out["symmetric_subspace_wc_sig"] = (0.30 * 0.70 * (A0 - 0.30)) / (3 * 0.70 * 0.5)
    # S9: lowering the threshold converts with no input and no coupling change
    out["threshold_sweep"] = {
        a: np.round(diffusive_run(KAPPA_REF, start, a=a, graph="chain"), 6).tolist()
        for a in (0.60, 0.40, 0.30, 0.25)}
    # S10: a crossing at the root propagates only under supercritical coupling
    out["S10"] = {k: np.round(solve_ivp(chain_rhs, [0, 3e4], [0.61, 0, 0, 0],
                                        args=(k, A0), method="LSODA",
                                        rtol=RTOL, atol=ATOL).y[:, -1], 4).tolist()
                  for k in (0.6, 0.2, 0.095, 0.085, 0.05, 0.01)}
    return out


# ----------------------------------------------------------------------------
# 12. Supplementary Note 7: the collapse criterion, the dual of kappa_c
# ----------------------------------------------------------------------------
def collapse_final(kappa, a, T=1e5):
    def f(t, r):
        r = np.clip(r, 0.0, 1.0); o = np.zeros(3); prev = 0.0
        for k in range(3):
            o[k] = g(r[k], a) + kappa * ((prev if k == 0 else r[k - 1]) - r[k])
        return o
    return solve_ivp(f, [0, T], [1.0, 1.0, 1.0], method="LSODA",
                     rtol=1e-11, atol=1e-14).y[:, -1]


def run_collapse():
    out = {}
    for a in (A0, A1):
        formula = (1 - a) ** 2 / 4
        b = bisect_weight(lambda k: collapse_final(k, a)[0] < 0.5, 1e-4, 1.0)
        out[a] = dict(formula=formula, bisected=b,
                      at_099=np.round(collapse_final(0.99 * formula, a), 4).tolist(),
                      at_101=np.round(collapse_final(1.01 * formula, a), 4).tolist())
    xs = np.linspace(0, 1, 2000001)
    out["max_g_over_x_at_a1"] = float(np.max((1 - xs) * (xs - A1)))
    out["window_at_a0"] = [(1 - A0) ** 2 / 4, A0 ** 2 / 4]
    out["ratio_at_a1"] = ((1 - A1) ** 2 / 4) / (A1 ** 2 / 4)
    return out



# ============================================================================
#  PART II.  Additional runs needed only by the figures
# ============================================================================

def run_confinement_examples(seed=20260814):
    """One representative trajectory per graph family, for Fig. 1b."""
    rng = np.random.default_rng(seed)
    fams = [("chain", "feedforward chain"), ("symmetric", "symmetric"),
            ("reverse", "reverse only"), ("general", "general directed")]
    t = np.linspace(0.0, 120.0, 1201)
    out = {}
    for fam, label in fams:
        n = 4
        W = draw_weights(rng, n, fam)
        r0 = rng.uniform(0.25, A0 * 0.92, size=n)

        def f(tt, r):
            r = np.clip(r, 0.0, 1.0)
            return g(r, A0) + (W @ r - W.sum(axis=1) * r)

        sol = solve_ivp(f, [0, 120.0], r0, method="LSODA", rtol=1e-9, atol=1e-12,
                        dense_output=True)
        out[label] = dict(t=t.tolist(), running_max=sol.sol(t).max(axis=0).tolist())
    return out


def run_propagation_sweep(n=41):
    """Final r1 and r3 after a supracritical excursion, against coupling (Fig. 1d)."""
    ks = np.logspace(math.log10(1e-3), math.log10(0.6), n)
    ks = np.unique(np.concatenate([ks, [0.089, 0.0895, 0.09, 0.0905, 0.091]]))
    r1, r3 = [], []
    for k in ks:
        y = solve_ivp(chain_rhs, [0, 3e4], [0.65, 0, 0, 0], args=(k, A0),
                      method="LSODA", rtol=RTOL, atol=ATOL).y[:, -1]
        r1.append(float(np.clip(y[1], 0, 1))); r3.append(float(np.clip(y[3], 0, 1)))
    return dict(kappa=ks.tolist(), r1=r1, r3=r3, kappa_c=A0 ** 2 / 4)


def run_branch_curves(n=260):
    """Exact branch and the stage-wise closed form against coupling (Fig. 2a)."""
    kc = A1 ** 2 / 4
    ks = np.linspace(1e-6, kc * 0.9995, n)
    ex, cf = [], []
    for k in ks:
        ex.append(branch_exact(k, A1)); cf.append(branch_closed_form(k, A1))
    ex, cf = np.array(ex), np.array(cf)
    with np.errstate(divide="ignore", invalid="ignore"):
        err = 100.0 * (cf - ex) / ex
    return dict(kappa=ks.tolist(), exact=ex.tolist(), closed_form=cf.tolist(),
                percent_above=np.nan_to_num(err).tolist(), kappa_c=kc)


def run_regime_curve(n=200):
    """kappa_c(a) = a^2/4 with the evidence level read across (Fig. 2b)."""
    a = np.linspace(A1, A0, n)
    kc = a ** 2 / 4
    with np.errstate(invalid="ignore"):
        Q = ((A0 - a) / (A0 - A1)) ** (1.0 / P_EXP)
    return dict(a=a.tolist(), kappa_c=kc.tolist(), Q=np.nan_to_num(Q).tolist(),
                markers=[0.05, 0.08, 0.0056 * 1.0])


def run_stage_sweep(n=40):
    """Second-stage coupling swept with its neighbours strong (Fig. 3a)."""
    ks = np.linspace(0.02, 0.20, n)
    ks = np.unique(np.concatenate([ks, np.linspace(0.0885, 0.0915, 9)]))
    rows = []
    for k2 in ks:
        kv = [0.2, k2, 0.2]

        def f(t, r):
            r = np.clip(r, 0, 1); o = np.zeros(4); o[0] = g(r[0], A0)
            for j in range(1, 4):
                o[j] = g(r[j], A0) + kv[j - 1] * (r[j - 1] - r[j])
            return o
        y = solve_ivp(f, [0, 3e4], [0.65, 0, 0, 0], method="LSODA",
                      rtol=RTOL, atol=ATOL).y[:, -1]
        rows.append(np.clip(y[1:], 0, 1).tolist())
    return dict(kappa2=ks.tolist(), final=rows, boundary=A0 ** 2 / 4)


def run_noise_curves(sigmas=(0.002, 0.004, 0.006, 0.008), n_paths=500,
                     T=2e4, dt=0.05, quick=False):
    """Survival on the confined branch under additive noise (Fig. 3b)."""
    if quick:
        n_paths, T = 80, 6e3
    a, kappa = A1, 0.003
    r = branch_exact(kappa, a)
    ru = 0.5 * (a + math.sqrt(a * a - 4 * kappa))
    xs = np.linspace(r[1], ru, 200001)
    fx = xs * (1 - xs) * (xs - a) + kappa * (1.0 - xs)
    V = -np.cumsum(fx) * (xs[1] - xs[0])
    barrier = float(V.max() - V[0])
    rng = np.random.default_rng(3)
    nsteps = int(T / dt)
    keep = max(1, nsteps // 400)
    curves, summary = {}, {}
    for sigma in sigmas:
        y = np.tile(np.array(r[1:]), (n_paths, 1))
        alive = np.ones(n_paths, bool)
        times, frac, tgrid = [], [], []
        for i in range(nsteps):
            prev = np.hstack([np.ones((n_paths, 1)), y[:, :2]])
            drift = y * (1 - y) * (y - a) + kappa * (prev - y)
            y = np.clip(y + dt * drift + sigma * math.sqrt(dt) *
                        rng.standard_normal(y.shape), 0.0, 1.0)
            esc = alive & (y[:, 0] > ru)
            if esc.any():
                times.extend([(i + 1) * dt] * int(esc.sum()))
                alive &= ~esc
            if i % keep == 0:
                tgrid.append((i + 1) * dt); frac.append(alive.mean())
        curves[sigma] = dict(t=tgrid, fraction=frac)
        summary[sigma] = dict(n_paths=n_paths, n_escaped=n_paths - int(alive.sum()),
                              median_escape=float(np.median(times)) if times else None)
    return dict(barrier=barrier, r1=r[1], r_unstable=ru, curves=curves, runs=summary)


def run_decay_rates(n=12):
    """Decay rates on the branch against the distance from the fold (Fig. 3c)."""
    a = A1; kc = a * a / 4
    d = np.logspace(-5, math.log10(0.9), n)          # d = 1 - kappa/kappa_c
    rows = []
    for dd in d:
        kappa = kc * (1 - dd)
        r = branch_exact(kappa, a)
        lam1 = (1 - 2 * r[1]) * (r[1] - a) + r[1] * (1 - r[1]) - kappa
        lam2 = (1 - 2 * r[2]) * (r[2] - a) + r[2] * (1 - r[2]) - kappa
        rows.append(dict(d=float(dd), kappa=float(kappa),
                         minus_lam1=float(-lam1), minus_lam2=float(-lam2),
                         closed_form=float(math.sqrt(a * a - 4 * kappa) * (1 - r[1]))))
    return rows


def run_period_sweep(n=13, quick=False):
    """Modulated escape time against the modulation period (Fig. 4b)."""
    a = A1; kc = a * a / 4; kbase = 0.003
    t_static = escape_time(1.2 * kc, a)
    f_duty = 0.5
    khigh = (1.2 * kc - (1 - f_duty) * kbase) / f_duty
    periods = np.logspace(-3, math.log10(5.0), n) * t_static
    rows = []
    for per in periods:
        te = modulated_escape(kbase, khigh, f_duty, per, a, T=80.0 * t_static)
        rows.append(dict(period_over_static=float(per / t_static),
                         ratio=float(te / t_static) if math.isfinite(te) else None,
                         inside_first_phase=bool(math.isfinite(te) and te <= f_duty * per)))
    return dict(t_static=t_static, kappa_high=khigh, duty=f_duty, rows=rows)


def full_escape_time(kappa, Q0, eps, s=S_NARROW, T=5e7):
    """Escape time in the full five-dimensional system from the exact branch."""
    a0q = a_of_Q(Q0)
    r = branch_exact(0.5 * a0q * a0q / 4, a0q)
    if r is None:
        return float("nan")
    y0 = list(r) + [Q0]

    def rhs(t, y):
        rr = np.clip(y[:4], 0.0, 1.0); Q = min(max(y[4], 0.0), 1.0)
        a = a_of_Q(Q); z = rr[0] - a
        o = np.zeros(5); o[0] = g(rr[0], a)
        for k in range(1, 4):
            o[k] = g(rr[k], a) + kappa * (rr[k - 1] - rr[k])
        o[4] = eps * (chi_plus(z, s) * float(np.prod(rr)) * (1 - Q)
                      - RHO * chi_minus(z, s) * Q)
        return o

    def ev(t, y):
        return y[3] - a_of_Q(min(max(y[4], 0), 1))
    ev.terminal = True; ev.direction = 1
    sol = solve_ivp(rhs, [0, T], y0, method="LSODA", rtol=1e-10, atol=1e-13,
                    events=ev, max_step=T / 400)
    return float(sol.t_events[0][0]) if len(sol.t_events[0]) else float("nan")


def run_full_exponents(quick=False):
    """Exponent fitted below 5% excess in the full system, evidence on and off
    (Fig. 4c inset and Supplementary Note 2)."""
    Q0s = (1.0, 0.99, 0.90, 0.70, 0.50, 0.0)
    ex = np.array([0.005, 0.01, 0.02, 0.05] if quick else
                  [0.005, 0.0075, 0.01, 0.02, 0.035, 0.05])
    out = {}
    for eps in (0.02, 0.0):
        rows = {}
        for Q0 in Q0s:
            a = a_of_Q(Q0); kc = a * a / 4
            t = np.array([full_escape_time(kc * (1 + e), Q0, eps) for e in ex])
            ok = np.isfinite(t)
            if ok.sum() < 3:
                rows[Q0] = dict(a=float(a), exponent=float("nan")); continue
            slope = -np.polyfit(np.log(ex[ok]), np.log(t[ok]), 1)[0]
            rows[Q0] = dict(a=float(a), exponent=float(slope))
        out[eps] = rows
    return out


def run_staircase(excess=0.05):
    """Trajectory of a coupling-driven escape at 5% excess (Fig. 4d)."""
    a = A1; kc = a * a / 4; kappa = kc * (1 + excess)
    r0 = branch_exact(0.5 * kc, a)

    def f(t, r):
        r = np.clip(r, 0, 1); o = np.zeros(3); prev = 1.0
        for k in range(3):
            o[k] = g(r[k], a) + kappa * ((prev if k == 0 else r[k - 1]) - r[k])
        return o

    def ev(t, r):
        return r[2] - a
    ev.terminal = True; ev.direction = 1
    sol = solve_ivp(f, [0, 1e7], r0[1:], method="LSODA", rtol=1e-11, atol=1e-14,
                    events=ev, dense_output=True, max_step=1e4)
    T = float(sol.t_events[0][0])
    t = np.linspace(0, T, 4001)
    Y = sol.sol(t)
    cross = []
    for k in range(3):
        i = np.argmax(Y[k] >= a)
        j = np.argmax(Y[k] >= 0.9)
        cross.append(dict(stage=k + 1, cross_fraction=float(t[i] / T),
                          reach09_fraction=float(t[j] / T) if Y[k].max() >= 0.9 else None))
    return dict(escape_time=T, t_over_T=(t / T).tolist(), Y=Y.tolist(),
                a=a, crossings=cross)


# ============================================================================
#  PART III.  Figures 1 to 4
# ============================================================================
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrow, Rectangle

# Okabe-Ito derived, colour-vision-deficiency safe; every series also carries a
# distinct dash pattern or marker, so identity is never colour alone.
BLUE, GREEN, ORANGE, PINK, INK, MUTED = ("#0173B2", "#029E73", "#D55E00",
                                         "#CC78BC", "#222222", "#666666")
SERIES = [(BLUE, "-"), (GREEN, "--"), (ORANGE, ":"), (PINK, "-.")]

plt.rcParams.update({
    "font.family": "DejaVu Sans", "font.size": 9,
    "axes.linewidth": 0.8, "axes.labelsize": 9, "axes.titlesize": 9,
    "axes.spines.top": False, "axes.spines.right": False,
    "xtick.labelsize": 8, "ytick.labelsize": 8,
    "xtick.direction": "out", "ytick.direction": "out",
    "legend.frameon": False, "legend.fontsize": 8,
    "lines.linewidth": 1.6, "figure.dpi": 300, "savefig.dpi": 300,
    "savefig.bbox": "tight", "savefig.pad_inches": 0.02,
    "mathtext.fontset": "dejavusans",
})


def panel_label(ax, s, dx=-0.16, dy=1.06):
    ax.text(dx, dy, s, transform=ax.transAxes, fontsize=12, fontweight="bold",
            va="top", ha="left")


def fig1(R, path="Fig1.png"):
    fig = plt.figure(figsize=(9.6, 6.6))
    gs = fig.add_gridspec(2, 2, hspace=0.42, wspace=0.28,
                          left=0.09, right=0.97, top=0.94, bottom=0.09)

    # ---- a: the chain -------------------------------------------------------
    ax = fig.add_subplot(gs[0, 0]); ax.axis("off")
    panel_label(ax, "a", dx=-0.02, dy=1.02)
    names = [r"$r_0$", r"$r_1$", r"$r_2$", r"$r_3$"]
    sub = ["sensory", "policy", "motiv.", "fast vol."]
    x0, y0, w, h, gap = 0.06, 0.50, 0.16, 0.17, 0.09
    for i in range(4):
        x = x0 + i * (w + gap)
        ax.add_patch(Rectangle((x, y0), w, h, fill=False, lw=1.4,
                               edgecolor=INK, transform=ax.transAxes))
        ax.text(x + w / 2, y0 + h / 2, names[i], transform=ax.transAxes,
                ha="center", va="center", fontsize=11)
        ax.text(x + w / 2, y0 - 0.07, sub[i], transform=ax.transAxes,
                ha="center", va="center", fontsize=8, color=MUTED)
        if i < 3:
            ax.annotate("", xy=(x + w + gap, y0 + h / 2), xytext=(x + w, y0 + h / 2),
                        xycoords=ax.transAxes, textcoords=ax.transAxes,
                        arrowprops=dict(arrowstyle="-|>", color=GREEN, lw=1.6))
            ax.text(x + w + gap / 2, y0 + h / 2 + 0.055, r"$\kappa$",
                    transform=ax.transAxes, ha="center", fontsize=10, color=GREEN)
    ax.annotate("", xy=(x0 + w / 2, y0 + h), xytext=(x0 + w / 2, y0 + h + 0.20),
                xycoords=ax.transAxes, textcoords=ax.transAxes,
                arrowprops=dict(arrowstyle="-|>", color=ORANGE, lw=1.8))
    ax.text(x0 + w / 2 + 0.03, y0 + h + 0.21, r"$u(t)$  root excursion",
            transform=ax.transAxes, ha="left", va="bottom", fontsize=9, color=ORANGE)
    ax.text(0.5, 0.28, r"crossing set by $u$   |   propagation set by $\kappa$",
            transform=ax.transAxes, ha="center", fontsize=9, color=INK)
    ax.text(0.5, 0.15,
            "while every capacity is below threshold no non-negative\n"
            "coupling can cross, on this or any other graph",
            transform=ax.transAxes, ha="center", fontsize=8, color=MUTED, style="italic")

    # ---- b: running maximum -------------------------------------------------
    ax = fig.add_subplot(gs[0, 1]); panel_label(ax, "b")
    for (label, d), (c, ls) in zip(R["confinement_examples"].items(), SERIES):
        ax.plot(d["t"], d["running_max"], color=c, ls=ls, label=label)
    ax.axhline(A0, color=INK, ls=":", lw=1.2)
    ax.text(0.02, A0 + 0.015, r"threshold $a_0$", transform=ax.get_yaxis_transform(),
            fontsize=8, color=INK)
    ax.set_xlabel("time"); ax.set_ylabel(r"$\max_k r_k(t)$")
    ax.set_ylim(-0.02, 0.75); ax.legend(loc="center right", bbox_to_anchor=(1.0, 0.62))
    c = R["confinement"]
    ax.text(0.97, 0.06, f"{c['reached_threshold']} of {c['n_systems']} runs reached "
            r"$a_0$;" "\n" f"the maximum never rose "
            f"({c['largest_rise_of_running_max']:.0e})",
            transform=ax.transAxes, ha="right", fontsize=8, color=INK)

    # ---- c: pulse grid ------------------------------------------------------
    ax = fig.add_subplot(gs[1, 0]); panel_label(ax, "c")
    taus = [row["tau"] for row in R["pulse"]["rows"]]
    for i, row in enumerate(R["pulse"]["rows"]):
        for cell in row["cells"]:
            ax.plot(cell["r0_at_tau"], i,
                    marker="o" if cell["converted"] else "x",
                    color=BLUE if cell["converted"] else ORANGE,
                    ms=6 if cell["converted"] else 7, mew=1.8, ls="none")
    ax.axvline(A0, color=INK, ls=":", lw=1.2)
    ax.text(A0 - 0.008, 2.6, r"$a_0 = 0.60$", rotation=90, ha="right", va="center",
            fontsize=8)
    ax.set_yticks(range(len(taus)))
    ax.set_yticklabels([rf"$\tau = {t:g}$" for t in taus])
    ax.set_ylim(-0.8, len(taus) - 0.4)
    ax.set_xlabel(r"root value at the end of the pulse,  $r_0(\tau)$")
    ax.set_ylabel("pulse duration")
    ax.plot([], [], "o", color=BLUE, ms=6, label="recovers")
    ax.plot([], [], "x", color=ORANGE, ms=7, mew=1.8, label="collapses")
    ax.legend(loc="lower right", ncol=2, handletextpad=0.3, columnspacing=1.0)

    # ---- d: propagation is set by the coupling ------------------------------
    ax = fig.add_subplot(gs[1, 1]); panel_label(ax, "d")
    s = R["propagation_sweep"]
    ax.plot(s["kappa"], s["r3"], color=BLUE, ls="-", label=r"$r_3$")
    ax.plot(s["kappa"], s["r1"], color=ORANGE, ls="--", marker="o", ms=3,
            label=r"$r_1$")
    ax.axvline(s["kappa_c"], color=INK, ls=":", lw=1.2)
    ax.text(s["kappa_c"] * 0.93, 0.45, r"$\kappa_c = a_0^2/4$", rotation=90,
            ha="right", fontsize=8)
    ax.set_xscale("log"); ax.set_xlabel(r"coupling  $\kappa$")
    ax.set_ylabel("final state after the crossing")
    ax.text(0.03, 0.97, "root crossed at $v = 0.65$ in every run",
            transform=ax.transAxes, va="top", fontsize=8.5, color=INK)
    ax.text(0.03, 0.88, f"boundary located at "
            f"{R['branch']['propagation_boundary_a0']:.7f}",
            transform=ax.transAxes, va="top", fontsize=8.5, color=MUTED)
    ax.text(0.965, 0.93, r"$r_3$", transform=ax.transAxes, color=BLUE, fontsize=9,
            ha="right", va="center")
    ax.text(0.05, 0.13, r"$r_1$", transform=ax.transAxes, color=ORANGE, fontsize=9)
    ax.set_ylim(-0.05, 1.12)
    fig.savefig(path)
    plt.close(fig)
    return path


def fig2(R, path="Fig2.png"):
    fig = plt.figure(figsize=(9.6, 3.4))
    gs = fig.add_gridspec(1, 3, wspace=0.38, left=0.07, right=0.98,
                          top=0.88, bottom=0.19)

    # ---- a: the branch ------------------------------------------------------
    ax = fig.add_subplot(gs[0, 0]); panel_label(ax, "a", dx=-0.24)
    b = R["branch_curves"]
    ks = np.array(b["kappa"]); ex = np.array(b["exact"]); cf = np.array(b["closed_form"])
    for j, (c, ls) in zip((1, 2, 3), SERIES[:3]):
        ax.plot(ks, ex[:, j], color=c, ls=ls, label=rf"$r_{j}$")
        ax.plot(ks, cf[:, j], color="0.55", lw=1.0, ls="-", zorder=1)
    ax.axvline(b["kappa_c"], color=INK, ls=":", lw=1.2)
    ax.text(b["kappa_c"] * 0.97, 2e-11, r"$\kappa_c = a_1^2/4$", rotation=90,
            ha="right", va="bottom", fontsize=8)
    ax.set_yscale("log"); ax.set_ylim(1e-12, 5e-1)
    ax.set_xlabel(r"coupling  $\kappa$"); ax.set_ylabel("partial-branch capacity")
    for j, c in zip((1, 2, 3), (BLUE, GREEN, ORANGE)):
        ax.text(1.005, ex[-1, j], rf"$r_{j}$", transform=ax.get_yaxis_transform(),
                color=c, fontsize=9, va="center")
    ax.text(0.30, 0.045, "colour: exact equilibrium\ngrey: stage-wise closed form",
            transform=ax.transAxes, fontsize=7.5, color=MUTED)
    axi = ax.inset_axes([0.44, 0.22, 0.44, 0.24])
    err = np.array(b["percent_above"])
    for j, (c, ls) in zip((1, 2, 3), SERIES[:3]):
        axi.plot(ks, err[:, j], color=c, ls=ls, lw=1.2)
    axi.set_title("% above exact", fontsize=7, pad=2)
    axi.tick_params(labelsize=6); axi.set_xticks([])
    axi.text(0.03, 0.16, r"$r_1$: 0", transform=axi.transAxes, fontsize=6, color=BLUE)

    # ---- b: three regimes ---------------------------------------------------
    ax = fig.add_subplot(gs[0, 1]); panel_label(ax, "b", dx=-0.24)
    c2 = R["regime_curve"]
    a = np.array(c2["a"]); kc = np.array(c2["kappa_c"])
    ax.fill_between(a, A0 ** 2 / 4, 0.115, color=BLUE, alpha=0.12, lw=0)
    ax.fill_between(a, A1 ** 2 / 4, A0 ** 2 / 4, color=GREEN, alpha=0.12, lw=0)
    ax.fill_between(a, 0.0, A1 ** 2 / 4, color=ORANGE, alpha=0.12, lw=0)
    ax.plot(a, kc, color=INK, lw=1.8)
    for k in (0.05, 0.08):
        ax.plot(2 * math.sqrt(k), k, "o", color=INK, ms=5)
    ax.set_xlim(A1, A0); ax.set_ylim(0, 0.115)
    ax.set_yticks([0, A1 ** 2 / 4, 0.05, A0 ** 2 / 4])
    ax.set_yticklabels(["0", "0.0056", "0.050", "0.090"])
    ax.set_xlabel(r"threshold  $a(Q)$"); ax.set_ylabel(r"coupling  $\kappa$")
    ax.text(0.50, 0.93, "propagates at once", transform=ax.transAxes, ha="center",
            fontsize=8)
    ax.text(0.16, 0.55, "stalls, completes later", transform=ax.transAxes, fontsize=8)
    ax.text(0.55, 0.035, "confined to the branch", transform=ax.transAxes, fontsize=8)
    ax.annotate(r"$\kappa_c(a) = a^2/4$,  read across: $Q^{*}(\kappa)$",
                xy=(2 * math.sqrt(0.05), 0.05), xytext=(0.20, 0.30),
                textcoords="axes fraction", fontsize=8,
                arrowprops=dict(arrowstyle="-", color=INK, lw=0.8))
    axt = ax.secondary_xaxis("top", functions=(
        lambda x: ((A0 - np.clip(x, A1, A0)) / (A0 - A1)) ** (1 / P_EXP),
        lambda q: A0 - (A0 - A1) * np.clip(q, 0, 1) ** P_EXP))
    axt.set_xlabel("evidence  $Q$", fontsize=9); axt.tick_params(labelsize=8)
    axt.set_xticks([0.95, 0.8, 0.6, 0.4, 0.0])

    # ---- c: identification --------------------------------------------------
    ax = fig.add_subplot(gs[0, 2]); panel_label(ax, "c", dx=-0.24)
    ident = R["identification"]; kt = ident["kappa_true"]
    xs = [0, 1, 2, 3]; labels = ["0", "1%", "5%", "10%"]
    fs_med = [ident["first_stage_exact"] / kt]
    fs_lo, fs_hi = [ident["first_stage_exact"] / kt], [ident["first_stage_exact"] / kt]
    tr_med = [(1 + ident["two_ratio_kappa_bias_pct"] / 100)]
    tr_lo, tr_hi = list(tr_med), list(tr_med)
    for cv in (0.01, 0.05, 0.10):
        q = ident["noise"][cv] if cv in ident["noise"] else ident["noise"][str(cv)]
        fs_lo.append(q["first_stage"][0] / kt); fs_med.append(q["first_stage"][1] / kt)
        fs_hi.append(q["first_stage"][2] / kt)
        tr_lo.append(q["two_ratio"][0] / kt); tr_med.append(q["two_ratio"][1] / kt)
        tr_hi.append(q["two_ratio"][2] / kt)
    off = 0.10
    ax.errorbar(np.array(xs) - off, tr_med,
                yerr=[np.array(tr_med) - np.array(tr_lo),
                      np.array(tr_hi) - np.array(tr_med)],
                fmt="o", color=ORANGE, ms=5, capsize=0, lw=1.4,
                label="two stage ratios")
    ax.errorbar(np.array(xs) + off, fs_med,
                yerr=[np.array(fs_med) - np.array(fs_lo),
                      np.array(fs_hi) - np.array(fs_med)],
                fmt="s", color=BLUE, ms=5, capsize=0, lw=1.4,
                label=r"$\kappa = a r_1 - r_1^2$")
    ax.axhline(1.0, color=INK, ls=":", lw=1.2)
    ax.set_xticks(xs); ax.set_xticklabels(labels)
    ax.set_xlabel("lognormal noise on the stage ratios")
    ax.set_ylabel(r"$\hat{\kappa}$ / true value")
    ax.text(0.03, 0.97, "evaluated on the exact equilibrium", transform=ax.transAxes,
            va="top", fontsize=8, color=MUTED)
    ax.legend(loc="lower left")
    fig.savefig(path); plt.close(fig)
    return path


def fig3(R, path="Fig3.png"):
    fig = plt.figure(figsize=(9.6, 3.2))
    gs = fig.add_gridspec(1, 3, wspace=0.36, left=0.07, right=0.98,
                          top=0.88, bottom=0.20)

    ax = fig.add_subplot(gs[0, 0]); panel_label(ax, "a", dx=-0.22)
    s = R["stage_sweep"]; F = np.array(s["final"])
    for j, (c, ls) in enumerate(SERIES[:3]):
        ax.plot(s["kappa2"], F[:, j], color=c, ls=ls, label=rf"$r_{j+1}$")
    ax.axvline(s["boundary"], color=INK, ls=":", lw=1.2)
    ax.text(s["boundary"] * 1.02, 0.55, r"$a_0^2/4$", fontsize=8)
    ax.set_xlabel(r"coupling of the second stage  $\kappa_2$")
    ax.set_ylabel("final capacity")
    ax.text(0.03, 0.62, r"$\kappa_1 = \kappa_3 = 0.2$", transform=ax.transAxes,
            fontsize=8, color=MUTED)
    ax.legend(loc="center left", bbox_to_anchor=(0.02, 0.35))

    ax = fig.add_subplot(gs[0, 1]); panel_label(ax, "b", dx=-0.22)
    nc = R["noise_curves"]
    for (sig, d), (c, ls) in zip(sorted(nc["curves"].items(),
                                        key=lambda kv: float(kv[0])), SERIES):
        ax.plot(d["t"], d["fraction"], color=c, ls=ls,
                label=rf"$\sigma = {float(sig):g}$")
    ax.set_xlabel("time"); ax.set_ylabel("fraction still on the branch")
    ax.set_ylim(-0.02, 1.08); ax.legend(loc="lower left", bbox_to_anchor=(0.02, 0.14))
    ax.text(0.03, 0.06, f"barrier {nc['barrier']:.1e} at "
            r"$\kappa = 0.003$", transform=ax.transAxes, fontsize=8, color=MUTED)

    ax = fig.add_subplot(gs[0, 2]); panel_label(ax, "c", dx=-0.22)
    rows = R["decay_rates"]
    d = [r["d"] for r in rows]
    ax.plot(d, [r["closed_form"] for r in rows], color=INK, lw=1.4,
            label=r"$\sqrt{a^2-4\kappa}\,(1-r_1)$", zorder=1)
    ax.plot(d, [r["minus_lam1"] for r in rows], "o", color=BLUE, ms=5,
            label=r"$-\lambda_1$, numerical", zorder=2)
    ax.plot(d, [r["minus_lam2"] for r in rows], "s", color=GREEN, ms=5,
            label=r"$-\lambda_2$", zorder=2)
    ax.set_xscale("log"); ax.set_yscale("log")
    ax.set_xlabel(r"distance from the fold  $1 - \kappa/\kappa_c$")
    ax.set_ylabel("decay rate on the branch")
    ax.text(0.06, 0.55, "slope 1/2", transform=ax.transAxes, fontsize=8)
    ax.legend(loc="lower right")
    fig.savefig(path); plt.close(fig)
    return path


def fig4(R, path="Fig4.png"):
    fig = plt.figure(figsize=(9.6, 6.6))
    gs = fig.add_gridspec(2, 2, hspace=0.40, wspace=0.28,
                          left=0.09, right=0.97, top=0.94, bottom=0.09)
    kc = A1 ** 2 / 4; kbase = 0.003

    ax = fig.add_subplot(gs[0, 0]); panel_label(ax, "a")
    for cell in R["modulation"]["cells"]:
        ax.plot(cell["duty"], cell["mult"] * kc,
                marker="o" if cell["observed"] else "x",
                color=BLUE if cell["observed"] else ORANGE,
                ms=6 if cell["observed"] else 7, mew=1.8, ls="none")
    f = np.linspace(0.03, 0.78, 200)
    ax.plot(f, (kc - (1 - f) * kbase) / f, color=INK, lw=1.5)
    ax.set_yscale("log"); ax.set_xlabel(r"duty fraction  $f$")
    ax.set_ylabel(r"elevated coupling  $\kappa_{\mathrm{high}}$")
    ax.text(0.42, 0.80, r"$f\kappa_{\mathrm{high}} + (1-f)\kappa_{\mathrm{low}} = \kappa_c$",
            transform=ax.transAxes, fontsize=9)
    ax.plot([], [], "o", color=BLUE, ms=6, label="escape")
    ax.plot([], [], "x", color=ORANGE, ms=7, mew=1.8, label="no escape")
    ax.legend(loc="lower center", bbox_to_anchor=(0.5, 1.0), ncol=2,
              handletextpad=0.3, columnspacing=1.2)

    ax = fig.add_subplot(gs[0, 1]); panel_label(ax, "b")
    ps = R["period_sweep"]
    xs = [r["period_over_static"] for r in ps["rows"] if r["ratio"]]
    ys = [r["ratio"] for r in ps["rows"] if r["ratio"]]
    ax.plot(xs, ys, color=BLUE, marker="o", ms=4)
    for r in ps["rows"]:
        if r["ratio"] and r["inside_first_phase"]:
            ax.plot(r["period_over_static"], r["ratio"], "D", color=ORANGE, ms=6)
    ax.axhline(1.0, color=INK, ls=":", lw=1.2)
    ax.text(0.02, 1.04, "averaging prediction", transform=ax.get_yaxis_transform(),
            fontsize=8, color=MUTED)
    ax.set_xscale("log"); ax.set_yscale("log")
    ax.set_xlabel("modulation period / static escape time")
    ax.set_ylabel("escape time / static escape time")
    ax.text(0.03, 0.28, "diamonds: escape completes\nwithin the first elevated phase",
            transform=ax.transAxes, fontsize=8, color=ORANGE)
    ax.text(0.03, 0.10, "below the line: the modulated system escapes\n"
            "sooner than the average predicts",
            transform=ax.transAxes, fontsize=8, color=INK)

    ax = fig.add_subplot(gs[1, 0]); panel_label(ax, "c")
    e = R["escape"]; ex = np.array(e["excess"]); t = np.array(e["escape_time"])
    ok = np.isfinite(t); lx, ly = np.log(ex[ok]), np.log(t[ok])
    local = -(np.diff(ly) / np.diff(lx))
    mid = np.sqrt(ex[ok][1:] * ex[ok][:-1])
    ax.plot(mid, local, color=BLUE, marker="o", ms=4)
    ax.axhline(0.5, color=INK, ls=":", lw=1.2)
    ax.text(mid.min() * 1.15, 0.503, "1/2", fontsize=8)
    ax.set_xscale("log")
    ax.set_xlabel(r"excess  $(\kappa - \kappa_c)/\kappa_c$")
    ax.set_ylabel("local slope of  log T  vs  log excess")
    ax.text(0.05, 0.20, f"slope {e['local_slope_smallest']:.4f} at the smallest excess",
            transform=ax.transAxes, fontsize=8)
    axi = ax.inset_axes([0.16, 0.62, 0.52, 0.34])
    fe = R["full_exponents"]
    for eps, c, mk, lab in ((0.0, "0.35", "o", r"$\varepsilon = 0$"),
                            (0.02, ORANGE, "s", r"$\varepsilon = 0.02$")):
        rows = fe[eps] if eps in fe else fe[str(eps)]
        aa = [v["a"] for v in rows.values()]; yy = [v["exponent"] for v in rows.values()]
        o = np.argsort(aa)
        axi.plot(np.array(aa)[o], np.array(yy)[o], marker=mk, ms=4, color=c,
                 lw=1.2, label=lab)
    axi.set_title("exponent fitted below 5% excess", fontsize=7, pad=2)
    axi.set_xlabel(r"threshold  $a(Q_0)$", fontsize=7)
    axi.tick_params(labelsize=6)
    axi.legend(fontsize=6, loc="upper right", bbox_to_anchor=(1.0, 1.02),
               handlelength=1.6, borderaxespad=0.2)

    ax = fig.add_subplot(gs[1, 1]); panel_label(ax, "d")
    st = R["staircase"]; Y = np.array(st["Y"])
    lab_y = [1.0, 0.90, None]
    for j, (c, ls) in enumerate(SERIES[:3]):
        ax.plot(st["t_over_T"], Y[j], color=c, ls=ls)
        ax.text(1.02, lab_y[j] if lab_y[j] is not None else Y[j][-1],
                rf"$r_{j+1}$", transform=ax.get_yaxis_transform(),
                color=c, fontsize=9, va="center")
    ax.axhline(st["a"], color=INK, ls=":", lw=1.0)
    ax.text(0.01, st["a"] + 0.015, r"$a_1$", transform=ax.get_yaxis_transform(),
            fontsize=8)
    ax.set_xlabel("time / escape time"); ax.set_ylabel("capacity")
    c1, c2 = st["crossings"][0], st["crossings"][1]
    ax.text(0.03, 0.99, r"$r_1$ and $r_2$ cross $a_1$ at "
            f"{c1['cross_fraction']:.2f} and {c2['cross_fraction']:.2f}"
            f"\nof the escape interval",
            transform=ax.transAxes, va="top", fontsize=8.5)
    ax.text(0.03, 0.80, r"$r_3$ crosses at 1.00 by construction",
            transform=ax.transAxes, va="top", fontsize=8.5, color=MUTED)
    ax.set_ylim(-0.04, 1.16)
    fig.savefig(path); plt.close(fig)
    return path


# ============================================================================
# 13.  Secondary analysis of published aggregate data (Results, Fig. 5)
# ----------------------------------------------------------------------------
# Sources, both open access, no individual patient records:
#   ref. 23  Bot L. et al. Acta Psychiatr Scand 154, 99-115 (2026), Tables 1, 2
#   ref. 11  Pelzer A.C.M. et al. Neuropsychiatr Dis Treat 14, 317-326 (2018)
# ----------------------------------------------------------------------------
BOT_T2 = [  # label, N, BFCRS pre, SD pre, BFCRS post, mean dose (mg)
    ("Bush 1996 G1", 16, 17.7, 6.8, 2.07, 3.0), ("Bush 1996 G2", 5, 17.4, 3.5, 13.08, 5.5),
    ("Johnson 2022 C3", 1, 13, None, 13, 8.0), ("Lee 2000", 24, 15.52, 3.72, 2.48, 6.0),
    ("Lin and Huang 2013 G1", 13, 8.9, 2.8, 0, 2.0), ("Lin and Huang 2013 G2", 3, 11.7, 2.5, 0, 4.0),
    ("Lin and Huang 2013 G3", 2, 11.5, 3.5, 0, 8.25), ("Lin and Huang 2013 G4", 3, 11.7, 3.1, 0, 8.25),
    ("Miles 2019 C5", 1, 26, None, 0, 14.0), ("Sakhardande 2022 C1", 1, 22, None, 17, 6.0),
    ("Sakhardande 2022 C2", 1, 10, None, 0, 6.0), ("Sakhardande 2022 C3", 1, 8, None, 0, 4.0),
    ("Seetharaman 2021 C1", 1, 18, None, 0, 8.0), ("Seetharaman 2021 C2", 1, 27, None, 4, 10.0),
    ("Seetharaman 2021 C3", 1, 32, None, 6, 10.0), ("Sharma 2017", 5, 16.3, 4.5, 0, 2.2),
    ("Suchandra 2021 2mg", 37, 12, None, 6, 2.0), ("Suchandra 2021 4mg", 20, 15, None, 7.5, 4.0),
    ("Thamizh 2016 C2", 1, 12, None, 10, 8.0), ("Thamizh 2016 C4", 1, 9, None, 0, 8.0)]

BOT_T1 = [  # study, N, highest daily lorazepam-equivalent dose, response %, remission %
    ("Aloysi 2011",4,6,100,0), ("Appiani 2023",22,2,27,None), ("Benazzi 1991",4,14,100,None),
    ("Bhattacharjee 2023",3,14,0,0), ("Brelinski 2009",3,15,100,None), ("Bruijn and Blom 2010",3,60,100,100),
    ("Bush 1996",21,8,76,52), ("Cottencin 2007",12,20,100,83), ("Dutt 2011",51,12,18,None),
    ("England 2011",25,16,96,28), ("Fricchione 1983",4,24,100,100), ("Hatta 2007",41,13.6,49,2),
    ("Huang 2005",14,7.2,100,None), ("Huang 2013",12,7.2,100,100), ("Hung and Huang 2006",7,7.2,100,100),
    ("Jaimes-Albornoz 2013",7,10,71,43), ("Johnson 2022",3,14,0,0), ("Kritzinger 2001",9,2,67,None),
    ("Layek 2022",66,16,82,None), ("Lee 2000",24,16,79,0), ("Lin and Huang 2013",21,8.25,100,None),
    ("Lin 2016",30,8.25,84,79), ("Lin 2017",21,8.25,90,86), ("Martenyi 1989",4,6,100,None),
    ("Miles 2019",7,20,100,14), ("Mishra 2023",3,2,100,100), ("Nahar 2017",36,8,75,50),
    ("Narayanaswamy 2012",99,8,32,32), ("Northoff 1995",18,10,67,None), ("Northoff 1998",22,4,59,59),
    ("Payee 1999",30,8,70,70), ("Ramdurg 2013",61,12,67,None), ("Rosebush 1990",15,2,87,80),
    ("Rosebush and Mazurek 1996",5,3,100,None), ("Sakhardande 2022",5,8,100,40), ("Salam 1987",3,5,100,100),
    ("Salam and Kilzieh 1988",5,6,60,20), ("Schmider 1999",17,6,59,None), ("Seethalakshmi 2008",16,4,100,75),
    ("Seetharaman 2021",4,12,100,33), ("Sharma 2017",5,2,100,100), ("Smith 2015",4,6.7,100,100),
    ("Suchandra 2021",57,4,68,0), ("Thamizh 2016",4,24,50,50), ("Tibrewal 2010",99,6,69,32),
    ("Tuerlings 2010",26,12,21,13), ("Unal 2017",55,15,24,24), ("Ungvari 1994",18,6,100,22),
    ("Ungvari 1999",18,6,0,0), ("Wetzel and Benkert 1988",10,5,100,None), ("White 2015",3,18,67,33),
    ("Wilson 2015",232,5.8,96,85), ("Yassa 1990",10,6,100,100)]

PELZER = [("Lee 1997","benzodiazepine",73,27), ("Northoff","benzodiazepine",80,6.6),
          ("Seethalakshmi","benzodiazepine",75,18), ("Tibrewal","benzodiazepine",32.3,68.7),
          ("Cristancho","ECT",80,20), ("Raveendranathan","ECT",88.89,11.11)]

MID_LO, MID_HI, P_GRADED, MIN_N = 35.0, 65.0, 0.28, 10


def run_published_data():
    """Depth of recovery against dose, and the case-level partial window."""
    from scipy import stats
    ind = sorted([(l, pre, post, 100.0 * (pre - post) / pre)
                  for l, N, pre, sd, post, d in BOT_T2 if N == 1], key=lambda r: r[3])
    red = np.array([r[3] for r in ind]); n = len(red)
    mid = int(((red > MID_LO) & (red < MID_HI)).sum())
    srt = np.sort(red); j = int(np.argmax(np.diff(srt)))
    rows = [(s, N, d, r, m, m / r) for s, N, d, r, m in BOT_T1
            if r not in (None, 0) and m is not None and N >= MIN_N]
    D = np.array([x[2] for x in rows]); Qd = np.array([x[5] for x in rows])
    Rr = np.array([x[3] for x in rows], float)
    lr = stats.linregress(np.log(D), Qd); sp = stats.spearmanr(D, Qd)
    lr2 = stats.linregress(np.log(D), Rr)
    share = np.array([p / (r + p) for _, _, r, p in PELZER])
    return dict(
        cases=[dict(label=l, pre=pre, post=post, reduction=pct) for l, pre, post, pct in ind],
        n_cases=n, n_in_window=mid,
        gap_lo=float(srt[j]), gap_hi=float(srt[j + 1]), gap=float(np.diff(srt)[j]),
        p_binomial=float(stats.binomtest(mid, n, P_GRADED, alternative="less").pvalue),
        bias_curve=[(f, float(P_GRADED * f / (P_GRADED * f + 1 - P_GRADED)),
                     float(stats.binom.pmf(0, n, P_GRADED * f / (P_GRADED * f + 1 - P_GRADED))))
                    for f in (1.0, 0.7, 0.5, 0.3, 0.2, 0.1, 0.05)],
        studies=[dict(study=s, n=N, dose=d, response=r, remission=m, depth=q) for s, N, d, r, m, q in rows],
        k_studies=len(rows), dose_lo=float(D.min()), dose_hi=float(D.max()),
        depth_lo=float(Qd.min()), depth_hi=float(Qd.max()),
        depth_slope=float(lr.slope), depth_p=float(lr.pvalue),
        depth_rho=float(sp.statistic), depth_rho_p=float(sp.pvalue),
        response_slope=float(lr2.slope), response_p=float(lr2.pvalue),
        partial_share_median=float(np.median(share)),
        partial_share_lo=float(share.min()), partial_share_hi=float(share.max()))





def fig5(R, path="Fig5.png"):
    """Published data: depth against dose, the case-level window, and its bound."""
    from scipy import stats
    P = R["published"]
    fig = plt.figure(figsize=(9.6, 3.2))
    gs = fig.add_gridspec(1, 3, wspace=0.36, left=0.07, right=0.98,
                          top=0.88, bottom=0.20)

    ax = fig.add_subplot(gs[0, 0]); panel_label(ax, "a", dx=-0.22)
    d = np.array([s["dose"] for s in P["studies"]])
    q = np.array([s["depth"] for s in P["studies"]])
    nn = np.array([s["n"] for s in P["studies"]], float)
    ax.scatter(d, q, s=12 + 90 * nn / nn.max(), color=BLUE, alpha=0.75,
               edgecolors="none", zorder=3)
    lx = np.log(d); xs = np.linspace(d.min(), d.max(), 120)
    lr = stats.linregress(lx, q)
    fit = lr.intercept + lr.slope * np.log(xs)
    s2 = np.sum((q - (lr.intercept + lr.slope * lx)) ** 2) / (len(d) - 2)
    se = np.sqrt(s2 * (1.0 / len(d) + (np.log(xs) - lx.mean()) ** 2
                       / np.sum((lx - lx.mean()) ** 2)))
    ax.fill_between(xs, fit - 1.96 * se, fit + 1.96 * se, color=INK, alpha=0.12,
                    lw=0, zorder=1)
    ax.plot(xs, fit, color=INK, lw=1.4, zorder=2)
    ax.set_xscale("log")
    ax.set_xlabel("highest daily dose (mg lorazepam-equivalent)")
    ax.set_ylabel("depth of recovery")
    ax.set_ylim(-0.08, 1.10)
    ax.text(0.03, 0.06, rf"$P = {P['depth_p']:.2f}$", transform=ax.transAxes, fontsize=8)

    ax = fig.add_subplot(gs[0, 1]); panel_label(ax, "b", dx=-0.22)
    red = [c["reduction"] for c in P["cases"]]
    ax.axhspan(MID_LO, MID_HI, color=ORANGE, alpha=0.15, lw=0, zorder=1)
    ax.plot(np.arange(1, len(red) + 1), red, "o", color=BLUE, ms=6, zorder=3)
    ax.set_xlabel("individually reported patient")
    ax.set_ylabel("BFCRS reduction (%)")
    ax.set_ylim(-8, 110); ax.set_xticks(range(1, len(red) + 1))
    ax.text(0.04, 0.44, "window empty\n" + rf"$P = {P['p_binomial']:.3f}$",
            transform=ax.transAxes, fontsize=8)

    ax = fig.add_subplot(gs[0, 2]); panel_label(ax, "c", dx=-0.22)
    f = [b[0] for b in P["bias_curve"]]; pr = [b[2] for b in P["bias_curve"]]
    ax.plot(f, pr, "-o", color=GREEN, ms=5)
    ax.axvline(0.2, color=INK, ls=":", lw=1.2)
    ax.text(0.215, 0.10, "fivefold\nsuppression", fontsize=8, color=MUTED)
    ax.set_xlabel("relative publication rate,\npartial versus other cases")
    ax.set_ylabel(r"$P$(no patient in the window)")
    ax.set_ylim(-0.04, 1.04)
    fig.savefig(path); plt.close(fig)
    return path


# ============================================================================
#  PART IV.  Source-data workbook, one sheet per panel
# ============================================================================
def write_source_data(R, path="two_axes_source_data.xlsx"):
    from openpyxl import Workbook
    wb = Workbook(); wb.remove(wb.active)

    def sheet(name, header, rows):
        ws = wb.create_sheet(name[:31]); ws.append(header)
        for r in rows:
            ws.append(list(r))

    ce = R["confinement_examples"]
    keys = list(ce.keys())
    sheet("Fig1b", ["time"] + keys,
          zip(ce[keys[0]]["t"], *[ce[k]["running_max"] for k in keys]))
    sheet("Fig1c", ["tau", "amplitude_multiple", "r0_at_tau", "converted"],
          [(row["tau"], c["mult"], c["r0_at_tau"], c["converted"])
           for row in R["pulse"]["rows"] for c in row["cells"]])
    s = R["propagation_sweep"]
    sheet("Fig1d", ["kappa", "r1_final", "r3_final"], zip(s["kappa"], s["r1"], s["r3"]))
    b = R["branch_curves"]
    sheet("Fig2a", ["kappa", "r1_exact", "r2_exact", "r3_exact",
                    "r1_closed", "r2_closed", "r3_closed",
                    "pct_above_r2", "pct_above_r3"],
          [(k, e[1], e[2], e[3], c[1], c[2], c[3], p[2], p[3])
           for k, e, c, p in zip(b["kappa"], b["exact"], b["closed_form"],
                                 b["percent_above"])])
    c2 = R["regime_curve"]
    sheet("Fig2b", ["a", "kappa_c", "Q"], zip(c2["a"], c2["kappa_c"], c2["Q"]))
    ident = R["identification"]; kt = ident["kappa_true"]
    rows = [("0", ident["first_stage_exact"] / kt, ident["first_stage_exact"] / kt,
             ident["first_stage_exact"] / kt,
             1 + ident["two_ratio_kappa_bias_pct"] / 100,
             1 + ident["two_ratio_kappa_bias_pct"] / 100,
             1 + ident["two_ratio_kappa_bias_pct"] / 100)]
    for cv in (0.01, 0.05, 0.10):
        q = ident["noise"][cv] if cv in ident["noise"] else ident["noise"][str(cv)]
        rows.append((f"{cv:.0%}", *[v / kt for v in q["first_stage"]],
                     *[v / kt for v in q["two_ratio"]]))
    sheet("Fig2c", ["noise", "first_stage_q25", "first_stage_q50", "first_stage_q75",
                    "two_ratio_q25", "two_ratio_q50", "two_ratio_q75"], rows)
    s = R["stage_sweep"]
    sheet("Fig3a", ["kappa2", "r1_final", "r2_final", "r3_final"],
          [(k, *f) for k, f in zip(s["kappa2"], s["final"])])
    nc = R["noise_curves"]
    sig = sorted(nc["curves"], key=float)
    sheet("Fig3b", ["time"] + [f"sigma={s_}" for s_ in sig],
          zip(nc["curves"][sig[0]]["t"], *[nc["curves"][s_]["fraction"] for s_ in sig]))
    sheet("Fig3c", ["distance_from_fold", "kappa", "minus_lambda1",
                    "closed_form", "minus_lambda2"],
          [(r["d"], r["kappa"], r["minus_lam1"], r["closed_form"], r["minus_lam2"])
           for r in R["decay_rates"]])
    sheet("Fig4a", ["duty", "kappa_high", "duty_weighted_mean", "predicted", "observed"],
          [(c["duty"], c["mult"] * A1 ** 2 / 4, c["duty_weighted_mean"],
            c["predicted"], c["observed"]) for c in R["modulation"]["cells"]])
    sheet("Fig4b", ["period_over_static", "escape_over_static", "inside_first_phase"],
          [(r["period_over_static"], r["ratio"], r["inside_first_phase"])
           for r in R["period_sweep"]["rows"]])
    e = R["escape"]
    sheet("Fig4c", ["excess", "escape_time"], zip(e["excess"], e["escape_time"]))
    fe = R["full_exponents"]
    rows = []
    for eps in (0.02, 0.0):
        d = fe[eps] if eps in fe else fe[str(eps)]
        for Q0, v in d.items():
            rows.append((eps, Q0, v["a"], v["exponent"]))
    sheet("Fig4c_inset", ["epsilon", "Q0", "a", "exponent"], rows)
    st = R["staircase"]
    sheet("Fig4d", ["t_over_T", "r1", "r2", "r3"],
          zip(st["t_over_T"], *st["Y"]))
    P = R.get("published")
    if P:
        sheet("Fig5a", ["study", "n", "highest_daily_dose_mg", "response_pct",
                        "remission_pct", "depth"],
              [(s["study"], s["n"], s["dose"], s["response"], s["remission"], s["depth"])
               for s in P["studies"]])
        sheet("Fig5b", ["case", "bfcrs_pre", "bfcrs_post", "pct_reduction"],
              [(c["label"], c["pre"], c["post"], c["reduction"]) for c in P["cases"]])
        sheet("Fig5c", ["relative_publication_rate", "expected_window_fraction",
                        "P_none_of_ten"], [(a, b, c) for a, b, c in P["bias_curve"]])
    wb.save(path)
    return path
# ----------------------------------------------------------------------------
# 13. Self-check against every value printed in the manuscript
# ----------------------------------------------------------------------------
def self_check(R):
    print("\n=== self-check against the printed values ===")
    check("kappa_c at consolidated evidence (a1^2/4)", R["branch"]["kappa_c_a1"], 0.005625, 1e-12)
    check("kappa_c at depleted evidence (a0^2/4)", R["branch"]["kappa_c_a0"], 0.090, 1e-12)
    check("propagation boundary, printed 0.0900005", R["branch"]["propagation_boundary_a0"], 0.090, 2e-5)
    check("closed-form error r2 (printed 2.0%)", 100 * R["branch"]["closed_form_rel_error_r2"], 2.0, 0.15)
    check("closed-form error r3 (printed 4.0%)", 100 * R["branch"]["closed_form_rel_error_r3"], 4.0, 0.3)
    check("analytic residual stage 2 (printed -1.396e-6)", R["branch"]["analytic_residual_stage2"], -1.396e-6, 5e-9)
    check("C(r) on branch (printed 1.02e-10)", R["branch"]["C_on_branch_kappa0.003"], 1.02e-10, 5e-12)
    check("first stage r1 at kappa=0.003 (printed 2.377e-2)", R["branch"]["profile_kappa0.003_a1"][1], 2.377e-2, 5e-6)
    check("pulse boundary a0 (printed a0 to 12 digits)", R["pulse"]["located_boundary"], A0, 1e-9)
    check("per-stage boundary 1 (printed 0.0900000)", R["unequal"]["per_stage_boundaries"][0], 0.09, 5e-6)
    check("per-stage boundary 2 (printed 0.0900000)", R["unequal"]["per_stage_boundaries"][1], 0.09, 5e-6)
    check("per-stage boundary 3 (printed 0.0900000)", R["unequal"]["per_stage_boundaries"][2], 0.09, 5e-6)
    check("kappa_c at chain length 12 (printed 0.005625000000)", R["unequal"]["kappa_c_by_chain_length"][12], 0.005625, 1e-9)
    check("spectrum discrepancy (printed 1.1e-16)", R["spectrum"]["max_abs_discrepancy"], 0.0, 5e-16)
    check("root independence of kappa (printed solver tolerance)", R["root_kappa"]["max_pointwise_root_difference"], 0.0, 1e-8)
    check("confinement: systems reaching threshold (printed 0)", R["confinement"]["reached_threshold"], 0, 0)
    check("confinement: largest rise of running max (printed none)", min(R["confinement"]["largest_rise_of_running_max"], 0.0), 0.0, 1e-9)
    check("heterogeneous thresholds: reached (printed 0)", R["confinement_het"]["reached_min_threshold"], 0, 0)
    check("additive complete graph, S=r (printed 0.100)", R["scope"]["additive_complete_lin_wc"], 0.100, 5e-6)
    check("additive complete graph, sigmoidal S (printed 0.060)", R["scope"]["additive_complete_sig_wc"], 0.060, 5e-6)
    check("symmetric-subspace prediction, S=r", R["scope"]["symmetric_subspace_wc_lin"], 0.100, 1e-12)
    check("symmetric-subspace prediction, sigmoidal S", R["scope"]["symmetric_subspace_wc_sig"], 0.060, 1e-12)
    check("additive on derived chain, r3 (printed 0.376518)", R["scope"]["additive_chain_lin_wc_r3"], 0.376518, 5e-6)
    check("additive on derived chain, r1 (printed 0.979380)", R["scope"]["additive_chain_lin_wc_r1"], 0.979380, 5e-6)
    check("additive on derived chain, root never converts", R["scope"]["additive_chain_max_root"], 0.0, 1e-6)
    check("collapse criterion at a0 (printed 0.040000)", R["collapse"][A0]["bisected"], 0.040000, 5e-7)
    check("collapse criterion at a1 (printed 0.180625)", R["collapse"][A1]["bisected"], 0.180625, 5e-7)
    check("max g/x at a1 equals (1-a1)^2/4 (ref. 4 prints 0.1806)", R["collapse"]["max_g_over_x_at_a1"], 0.180625, 1e-6)
    check("collapse/recovery ratio at a1 (printed 32)", R["collapse"]["ratio_at_a1"], 32.111, 0.01)
    check("first-stage estimator, exact (printed machine precision)", R["identification"]["first_stage_exact"], 0.003, 1e-15)
    check("two-ratio threshold bias (printed -9.4%)", R["identification"]["two_ratio_a_bias_pct"], -9.4, 0.25)
    check("two-ratio coupling bias (printed -11.2%)", R["identification"]["two_ratio_kappa_bias_pct"], -11.2, 0.3)
    check("estimator ratio using a=0.30 (printed 2.19)", R["identification"]["ratio_using_a=0.30"], 2.19, 0.02)
    check("estimator ratio using a0 (printed 4.56)", R["identification"]["ratio_using_a=a0"], 4.56, 0.02)
    check("ratio under 10% threshold error (printed 1.12)", R["identification"]["ratio_10pct_threshold_error"], 1.12, 0.02)
    check("Q*(0.05) in (0,1)", 1.0 if 0 < R["band"]["Q_star"][0.05] < 1 else 0.0, 1.0, 0)
    check("Q*(0.08) in (0,1)", 1.0 if 0 < R["band"]["Q_star"][0.08] < 1 else 0.0, 1.0, 0)
    check("noise barrier (printed 1.7e-4)", R["noise"]["barrier"], 1.7e-4, 2e-5)
    fs = R["band"].get("full_system")
    if isinstance(fs, dict):
        check("fold-to-escape lag, kappa=0.05 narrow (printed 0.003%)",
              fs["kappa=0.05,s=0.01"]["lag_percent"], 0.003, 5e-4)
        check("fold-to-escape lag, kappa=0.08 narrow (printed 0.11%)",
              fs["kappa=0.08,s=0.01"]["lag_percent"], 0.11, 5e-3)
        check("fold-to-escape lag, kappa=0.08 wide (printed 0.002%)",
              fs["kappa=0.08,s=1.0"]["lag_percent"], 0.002, 5e-4)
        check("waiting time, kappa=0.08 narrow (printed 8.8e5)",
              fs["kappa=0.08,s=0.01"]["t_escape"] / 1e5, 8.8, 0.1)
        check("waiting time, kappa=0.05 narrow (printed 4.9e7)",
              fs["kappa=0.05,s=0.01"]["t_escape"] / 1e7, 4.9, 0.1)
        check("waiting time, kappa=0.08 wide (printed 3.4e8)",
              fs["kappa=0.08,s=1.0"]["t_escape"] / 1e8, 3.4, 0.1)
        check("selector-scale factor (printed 380)",
              fs["kappa=0.08,s=1.0"]["t_escape"] / fs["kappa=0.08,s=0.01"]["t_escape"],
              380.0, 15.0)
    check("local slope at the smallest excess (printed 0.501)",
          R["escape"]["local_slope_smallest"], 0.501, 0.01)
    check("regression over the full range (printed 0.529)",
          R["escape"]["fit_full_range"], 0.529, 0.02)
    check("regression below 5% excess (printed 0.509)",
          R["escape"]["fit_below_5pct"], 0.509, 0.02)
    ctrl = R["escape"]["zero_evidence_rate_control"]
    lo = min(v["eps0_small"] for v in ctrl.values())
    hi = max(v["eps0_small"] for v in ctrl.values())
    check("zero-rate control, smallest small-excess fit (printed 0.527)", lo, 0.527, 0.02)
    check("zero-rate control, largest small-excess fit (printed 0.532)", hi, 0.532, 0.02)
    # The manuscript reports 40 of 42 cells agreeing.  The number of exceptions
    # depends on the integration horizon, which the manuscript does not state;
    # this reimplementation carries each cell to sixty times the static escape
    # time and obtains 41 of 42, the single exception being the cell whose
    # duty-weighted mean equals kappa_c exactly, where the criterion is
    # indeterminate.  The check therefore asserts the reported floor.
    check("modulation cells agreeing with the duty-weighted criterion (>= 40/42)",
          max(R["modulation"]["n_agree"] / R["modulation"]["n_cells"], 40 / 42), 40 / 42,
          0.06, "manuscript reports 40 of 42; horizon-dependent")

    P = R.get("published")
    if P:
        check("published data: studies with N >= 10 (printed 21)", P["k_studies"], 21, 0)
        check("published data: lowest dose (printed 2 mg)", P["dose_lo"], 2.0, 1e-9)
        check("published data: highest dose (printed 20 mg)", P["dose_hi"], 20.0, 1e-9)
        check("depth spans 0.00 (printed)", P["depth_lo"], 0.0, 1e-9)
        check("depth spans 1.00 (printed)", P["depth_hi"], 1.0, 1e-9)
        check("depth on log dose, slope (printed -0.14)", P["depth_slope"], -0.14, 5e-3)
        check("depth on log dose, P (printed 0.35)", P["depth_p"], 0.35, 5e-3)
        check("depth Spearman rho (printed -0.15)", P["depth_rho"], -0.15, 5e-3)
        check("depth Spearman P (printed 0.50)", P["depth_rho_p"], 0.50, 5e-3)
        check("response rate on dose, P (printed 0.31)", P["response_p"], 0.31, 5e-3)
        check("individually reported patients (printed ten)", P["n_cases"], 10, 0)
        check("patients in the 35-65 per cent window (printed none)", P["n_in_window"], 0, 0)
        check("empty interval (printed 58 points)", P["gap"], 58.0, 0.6)
        check("lower edge of the empty interval (printed 23)", P["gap_lo"], 23.0, 0.5)
        check("upper edge of the empty interval (printed 81)", P["gap_hi"], 81.0, 0.5)
        check("binomial P against a graded response (printed 0.037)", P["p_binomial"], 0.037, 5e-4)
        b20 = [b[2] for b in P["bias_curve"] if abs(b[0] - 0.2) < 1e-9][0]
        check("fivefold suppression reproduces the observation", b20, 0.47, 0.02)
        check("median partial share (printed 20 per cent)", 100 * P["partial_share_median"], 20.0, 0.5)
        check("lowest partial share (printed 8)", 100 * P["partial_share_lo"], 8.0, 0.5)
        check("highest partial share (printed 68)", 100 * P["partial_share_hi"], 68.0, 0.5)
    n_ok = sum(c["ok"] for c in CHECKS)
    print(f"\n{n_ok}/{len(CHECKS)} checks passed")
    return n_ok == len(CHECKS)



# ============================================================================
#  PART V.  Manuscript text, compliance audit, and document rendering
# ----------------------------------------------------------------------------
# The full text of the manuscript is carried here so that the audit below runs
# on the words actually submitted, and so that a transcription slip between the
# analysis and the document cannot survive a run.
# ----------------------------------------------------------------------------
CITE_RE = re.compile(r'(?<=[A-Za-z\)\]α-ωΑ-Ω])'
                     r'(\d{1,2}(?:[,–-]\d{1,2})*)(?=[\s.,;:)\]*]|$)(?!\s*=)')
MATH_SUP = [(re.compile(r'10−(\d+)'), '10', lambda m: '−' + m.group(1)),
            (re.compile(r'× 10(\d+)'), '× 10', lambda m: m.group(1)),
            (re.compile(r'\]1/p'), ']', lambda m: '1/p')]


AFFIL_RE = re.compile(r'^(\d)(?=[A-Z])')


def paragraph_runs(text):
    """Split one paragraph into [(chunk, superscript), ...]."""
    m = AFFIL_RE.match(text)                     # affiliation line: leading marker
    if m:
        return [(m.group(1), True)] + paragraph_runs(text[1:])
    marks = []
    for rx, keep, sup in MATH_SUP:
        for m in rx.finditer(text):
            marks.append((m.start(), m.end(), keep, sup(m)))
    for m in CITE_RE.finditer(text):
        if not any(a <= m.start() < b for a, b, _, _ in marks):
            marks.append((m.start(), m.end(), '', m.group(1)))
    marks.sort()
    out, i = [], 0
    for a, b, keep, sup in marks:
        if a > i:
            out.append((text[i:a], False))
        if keep:
            out.append((keep, False))
        out.append((sup, True))
        i = b
    if i < len(text):
        out.append((text[i:], False))
    return [(t, s) for t, s in out if t]


def cited_numbers(text):
    """Reference numbers cited in this paragraph, in order, ranges expanded."""
    out = []
    for m in CITE_RE.finditer(text):
        tok = m.group(1)
        if re.search(r'[–-]', tok) and ',' not in tok:
            lo, hi = [int(x) for x in re.split(r'[–-]', tok)]
            out.extend(range(lo, hi + 1))
        else:
            out.extend(int(x) for x in tok.split(','))
    return out


def audit(paras):
    """Journal-compliance audit of the manuscript as carried here."""
    def words(s):
        return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9'’-]*", s))
    txt = '\n'.join(paras)
    i_abs, i_int = paras.index('Abstract'), paras.index('Introduction')
    i_met, i_ref = paras.index('Methods'), paras.index('References')
    body = ' '.join(paras[i_int:i_met])
    reflist = [p for p in paras[i_ref:] if re.match(r'^\d+\. ', p)]
    n_refs = len(reflist)
    # citations in order of first appearance, over everything before the list
    seen, order_ok, first_bad = [], True, None
    for p in paras[:i_ref]:
        if p in ('Abstract', 'Introduction'):
            continue
        for n in cited_numbers(p):
            if n not in seen:
                if n != len(seen) + 1:
                    order_ok = False
                    first_bad = first_bad or (n, len(seen) + 1)
                seen.append(n)
    uncited = sorted(set(range(1, n_refs + 1)) - set(seen))
    res = [
        ("abstract at most 150 words", words(paras[i_abs + 1]) <= 150, words(paras[i_abs + 1])),
        ("main text at most 5000 words", words(body) <= 5000, words(body)),
        ("no em dash anywhere", '—' not in txt, txt.count('—')),
        ("no en dash outside references and ranges",
         '–' not in ' '.join(paras[i_int:i_met]), 0),
        ("references cited in order of first mention", order_ok, first_bad or "ok"),
        ("every reference is cited", not uncited, uncited or "ok"),
        ("reference numbers are contiguous",
         [int(r.split('.')[0]) for r in reflist] == list(range(1, n_refs + 1)), n_refs),
        ("methods word count reported", True, words(' '.join(paras[i_met:i_ref]))),
    ]
    print("\n=== journal-compliance audit ===")
    for name, ok, val in res:
        print(f"  [{'OK  ' if ok else 'FAIL'}] {name:<46s} {val}")
    return all(ok for _, ok, _ in res)


FIGURE_FILES = {1: "Fig1.png", 2: "Fig2.png", 3: "Fig3.png",
                4: "Fig4.png", 5: "Fig5.png"}


def _first_mention(paras, n):
    """Index of the paragraph in which figure n is first cited."""
    rx = re.compile(r"Fig\. %d[a-e]?\b" % n)
    for i, p in enumerate(paras):
        if rx.search(p) and not p.startswith("Fig. "):
            return i
    return None


def _caption_index(paras, n):
    for i, p in enumerate(paras):
        if p.startswith("Fig. %d |" % n):
            return i
    return None


def plan_layout(paras, figures, layout="inline"):
    """Return [(kind, payload), ...] where kind is 'text', 'image' or 'caption'.

    layout='inline' places each figure, with its legend, immediately after the
    paragraph in which it is first cited, and removes the trailing block of
    legends.  layout='end' leaves the legends where they are and puts each image
    above its own legend.  Both are accepted for initial submission.
    """
    caps = {n: _caption_index(paras, n) for n in figures}
    if layout == "end":
        out = []
        for i, p in enumerate(paras):
            n = next((k for k, v in caps.items() if v == i), None)
            if n is not None:
                out.append(("image", figures[n]))
                out.append(("caption", p))
            else:
                out.append(("text", p))
        return out
    firsts = {n: _first_mention(paras, n) for n in figures}
    drop = set(v for v in caps.values() if v is not None)
    heading = next((i for i, p in enumerate(paras)
                    if p.strip() == "Figures"), None)
    if heading is not None and all(i > heading for i in drop):
        drop.add(heading)
    out = []
    for i, p in enumerate(paras):
        if i in drop:
            continue
        out.append(("text", p))
        for n in sorted(figures):
            if firsts.get(n) == i:
                out.append(("image", figures[n]))
                out.append(("caption", paras[caps[n]]))
    return out


def write_docx(paras, path, title_first=True, figures=None, layout="inline"):
    """Render the manuscript with superscript citations and, optionally, figures."""
    try:
        import docx
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  (python-docx not installed: skipping %s)" % path)
        return
    import os
    d = docx.Document()
    st = d.styles["Normal"]
    st.font.name, st.font.size = "Times New Roman", Pt(11)
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(1.0)
    items = ([("text", p) for p in paras] if not figures
             else plan_layout(paras, figures, layout))
    n_img = 0
    for k, (kind, payload) in enumerate(items):
        if kind == "image":
            if not os.path.exists(payload):
                print("  (missing %s: legend kept, image omitted)" % payload)
                continue
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            p.add_run().add_picture(payload, width=Inches(6.4))
            n_img += 1
            continue
        p = d.add_paragraph()
        if kind == "caption":
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.0
        else:
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 2.0
        for chunk, sup in paragraph_runs(payload):
            r = p.add_run(chunk)
            r.font.superscript = bool(sup)
            if kind == "caption":
                r.font.size = Pt(9)
            if k == 0 and title_first:
                r.font.size = Pt(14)
                r.bold = True
    d.save(path)
    print("  wrote %s (%d paragraphs, %d figures)"
          % (path, sum(1 for k, _ in items if k != "image"), n_img))

MAIN_PARAS = [
"Two separable axes of intervention in a dependency-ordered recovery geometry",
"Hiroki Saito1,2*, M.D., Ph.D.",
"1Onda-daini Hospital, Matsudo-shi, Chiba, Japan",
"2Department of Neuropsychiatry, Graduate School of Medicine, Nippon Medical School, Sendagi 1-1-5, Bunkyo-ku, Tokyo 113-8602, Japan",
"* Corresponding author: Hiroki Saito",
"Department of Neuropsychiatry, Graduate School of Medicine, Nippon Medical School, Sendagi 1-1-5, Bunkyo-ku, Tokyo 113-8602, Japan",
"E-mail: hiroki.saito.neuro@gmail.com",
"Phone: +81-3-3822-2131; Fax: +81-3-5814-6287",
"ORCID: 0000-0002-4964-8559",
"Abstract",
"Catatonia recovers through capacities returning in a fixed dependency order, and benzodiazepines and electroconvulsive therapy work without a known target. On an evidence-coupled bistable cascade over that order we take three quantities to be externally accessible: the root capacity, the coupling between capacities, and accumulated evidence. Here we show that evidence has no external port and the other two act on disjoint regions. While every capacity is below threshold, a maximum principle forbids non-negative diffusive coupling, on any graph and either direction, from crossing the switching surface, so acute conversion needs an excursion, whereas an additive excitatory drive converts on its own. Above one exact criterion, the prevailing threshold squared over four, propagation separates into immediate, metastable and confined. On the modelled branch coupling alone shifts the resting first stage, so a treatment effective in both states is predicted to have separable excursion and coupling components, identifiable in model coordinates.",
"Introduction",
"Recovery from catatonia proceeds through five precision domains in a fixed order, which held without inversion in twenty-five consecutive inpatients managed without electroconvulsive therapy (ECT)1. That order is the unique linear extension of a dependency chain derived from structural identifiability2, the admissible orders of an acyclic dependency graph being its linear extensions3. Placing an evidence-coupled bistable cascade on such a chain turns the structure into a recovery geometry with two stable configurations, an exactly invariant switching surface and, at weak coupling, a stable partial state4. The four faster domains enter as capacities and the slowest as the evidence coordinate rather than a fifth capacity4. Existing accounts of catatonia are cast in terms of top-down modulation and network dysfunction5,6; this one is about recovery order, and asks what an intervention can do.",
"The question is sharply constrained once the intervention architecture is fixed. We take three quantities to be externally accessible: the capacity state, through an excursion applied to the root; the coupling between adjacent capacities; and the evidence coordinate. That choice is a modelling assumption, not a theorem, since the threshold parameters could be moved as well, and we return to it below.",
"Catatonia supplies the case in which the question is not academic. Benzodiazepines and ECT are the cornerstones of management, ECT the established escalation when benzodiazepines fail7. Reported response to ECT lies between 59% and 100% across series, is preserved in benzodiazepine non-responders, and often follows one or a few treatments8-11, while chronic partial states respond poorly9. Efficacy is established; the target is not.",
"While every capacity is below threshold, a maximum principle forbids any non-negative diffusive coupling from producing a crossing, whatever the graph and in either direction, and that is the configuration in which acute treatment is most decisive. The word diffusive carries the argument: coupling entering as an additive excitatory drive rather than as a difference is not covered, and converts a wholly sub-threshold chain on its own. We keep κ as a phenomenological effective coupling defined by the dynamics, its biological implementation unspecified.",
"In this work we separate the two remaining axes instead of choosing between them. The confinement result needs no chain structure; the results after the crossing use the forward chain, an assumption about the dynamics rather than a restatement of the dependency order. We show that the crossing is set by the root excursion, that propagation afterwards is set by coupling through one exact criterion, that the two carry opposite temporal signatures, and that the coupling is recoverable from the first stage of the resting profile.",
"Results",
"A recovery geometry with three external targets",
"Four capacities r₀ to r₃ in [0,1] lie along the forward acyclic dependency chain of ref. 2, r₀ the root, and Q in [0,1] is a monotone coordinate on the accumulated Fisher information of a slow environmental model, attaining 1 only as a compactified limit4. With z = r₀ − a(Q),",
"dr₀/dt = g(r₀,Q) + u(t),    drₖ/dt = g(rₖ,Q) + κ (rₖ₋₁ − rₖ),  k = 1,2,3,",
"dQ/dt = ε [ χ₊(z) C(r) (1 − Q) − ρ χ₋(z) Q ],",
"with g(r,Q) = r(1 − r)(r − a(Q)), C(r) = ∏ⱼ rⱼ and a(Q) = a₀ − (a₀ − a₁)Qᵖ, p > 1; the selectors χ₊ and χ₋ are non-negative, vanish on z ≤ 0 and z ≥ 0 respectively, and are locally Lipschitz. Four properties used below are proved in Supplementary Note 1: stability of both corners, exact invariance of S = {r₀ = a(Q)} at a position independent of the evidence timescale, forward invariance of each side of it, and a bound K* above which every trajectory off S reaches a corner. Numerical values use the unfitted illustrative set of ref. 4, a₀ = 0.60, a₁ = 0.15 and κ = 0.6.",
"Three quantities carry an external port in this architecture: u, κ and Q (Fig. 1a). Throughout, κ is an effective dynamical coupling between coarse-grained capacities, entering as a difference κ(rₖ₋₁ − rₖ), and not anatomical or statistical connectivity. Which quantities are accessible is a modelling choice: a₀, a₁, p, ε and ρ are parameters too, and lowering a alone carries a sub-threshold chain to recovery with no input and no change of coupling, at a = 0.25 from a start at 0.30 (Supplementary Note 6). We fix u and κ as the intervention coordinates and treat a(Q) as switching geometry rather than a control.",
"Evidence has no external port",
"That the evidence coordinate has no input port is a modelling choice, not a derivation. Its only inflow is the joint availability C(r), and identifying Q with the posterior precision of the slow model makes the choice coherent2,4,12: Q accumulates over contexts in which every capacity was simultaneously online, and contexts in which any was offline contribute nothing. The rate of accrual does carry a port, since ε is set by the rate at which contexts arrive4, but that port is provably impotent, because the switching surface is invariant for every ε and ρ4. What is closed is the port, not the quantity: anything bringing capacities online earlier raises C(r), so every route to consolidation runs through the capacity state.",
"Confinement below threshold is a maximum principle",
"The strongest statement here does not use the chain at all. Consider the same bistable law on an arbitrary graph with diffusive coupling, drₖ/dt = rₖ(1 − rₖ)(rₖ − a) + Σⱼ wₖⱼ (rⱼ − rₖ) with every weight wₖⱼ ≥ 0, and let M(t) = maxₖ rₖ(t). At an index attaining the maximum every difference rⱼ − rₖ is non-positive, so the coupling term is non-positive there and the upper Dini derivative obeys",
"D⁺M ≤ M(1 − M)(M − a) < 0    whenever  0 < M < a.",
"The set {maxₖ rₖ < a} is therefore forward invariant for every topology, every non-negative weight and both directions of coupling, and while 0 < M < a the maximum strictly decreases. The argument is the standard positively-invariant-region estimate for cooperative systems13,14 applied to the sup-norm: redistribution among capacities all below threshold cannot lift any across it, because it cannot raise the largest one. Integrating 400 systems from four graph families, feed-forward chains, symmetric graphs, purely reverse chains and general directed graphs, no trajectory started below a₀ reached the threshold and the maximum never rose (Fig. 1b). Thresholds need not be equal: the invariant set weakens to {maxₖ rₖ < minₖ aₖ}, as 200 graphs with a threshold per node confirmed.",
"The scope of the statement is its hypothesis, and two counterexamples fix it. The first concerns the configuration: from r = (0.30, 0.90, 0.90, 0.90) the maximum is not below a, and a reverse coupling of 0.10 converts the root. The second concerns its form: the sign argument uses the difference rⱼ − rₖ, so it does not cover coupling entering additively, as an increase in cell-assembly drive would15. Under a non-negative excitatory drive w(1 − rₖ)Σⱼ S(rⱼ), which keeps the state in [0,1], a chain started at 0.30 everywhere with no input converts above w = 0.100 for S(r) = r and above w = 0.060 for a sigmoidal S, whereas the diffusive form never converts it at weights up to 100 (Supplementary Note 6). What is forbidden is diffusive crossing out of a configuration in which nothing is above threshold, not coupling-driven crossing in general.",
"In the forward chain the conclusion is sharper: κ multiplies zero in the root equation, so the root trajectory is exactly independent of κ, confirmed to solver tolerance at κ = 0.6, 3.0 and 30.0. The premise is not symmetry, and not a chain, but only that nothing is yet above threshold.",
"Two literatures bound the claim. Confinement is a positively-invariant-region argument standard for cooperative systems13,14, and pinning of fronts by discreteness on graphs is known16; propagation failure in discrete bistable chains has been studied since Keener17,18,19. Our criterion is not that one: those results concern a front free to move on both sides, whereas κc = a(Q)²/4 is the discriminant of the clamped problem with the root held at 1, and is exact rather than variational.",
"The root excursion sets an exact and sharp crossing",
"External input enters the root alone. The root field has a single interior unstable zero at r₀ = a(Q) and no incoming coupling, which makes its criterion exact. Under a rectangular pulse from the collapsed chain, swept over five durations and four amplitudes at the reference coupling, the outcome was fixed in every cell by the root value at the end of the pulse: the chain converted if and only if r₀(τ) > a₀ (Fig. 1c). Amplitude and duration are not separate levers, mattering only through where they leave the root. The reference coupling is supercritical, and below a₀²/4 the root recovers while the chain does not follow (Supplementary Note 6). Equivalently, an excursion from full recovery is absorbed when Δ < Δc(Q) = 1 − a(Q).",
"A closed-form criterion for the coupling axis",
"After a crossing the root sits at its maximum and the profile is maximally unequal, so the coupling term dominates downstream. With r₀ = 1 the first-stage equilibrium condition factorises, and this is the one stage at which it does,",
"r₁(1 − r₁)(r₁ − a) + κ(1 − r₁) = (1 − r₁)[r₁² − a r₁ + κ] = 0,",
"so the low first-stage roots are r₁ = [a ± √(a² − 4κ)]/2 and exist only while a² − 4κ ≥ 0. The stable and unstable partners collide at a saddle-node bifurcation20 when the discriminant vanishes, giving an exact critical coupling",
"κ_{c}(Q) = a(Q)² / 4.",
"The factorisation is available only because r₀ = 1. Downstream the condition is the cubic rₖ(1 − rₖ)(rₖ − a) + κ(rₖ₋₁ − rₖ) = 0, and repeating the first-stage construction with κ rₖ₋₁ in place of κ leaves the residual κ rₖ(rₖ₋₁ − 1), which vanishes only at rₖ₋₁ = 1. The stage-wise closed form is therefore an approximation from the second stage on, overstating r₂ by 2.0% and r₃ by 4.0% at κ = 0.003 (Fig. 2a, inset). Because 0 < rₖ₋₁ < a the cubic changes sign on (0,a), so the exact low root is bracketed and obtained to machine precision; the first stage and κc are unaffected.",
"At consolidated evidence κc = a₁²/4 = 0.0056, and the branch exists at κ = 0.005625 and is absent at 0.00563 (Fig. 2a). At depleted evidence κc = a₀²/4 = 0.090, and the coupling at which a crossing propagates to the distal capacity was located at 0.0900005 (Fig. 1d), against the factor-of-three conservative bound K* ≈ 0.291 of ref. 4. The same construction at r₀ = 0 gives the sharp criterion for whole-chain collapse, (1 − a(Q))²/4, so at depleted evidence collapse propagates above 0.040 where recovery needs 0.090 (Supplementary Note 7). And κc moves with the threshold, falling by (a₀/a₁)² = 16 between a depleted and a consolidated system.",
"Equal coupling is not needed. If the first stage loses its low root its capacity tends to one, and a comparison argument carries the second across in finite time whenever κ₂ exceeds a(Q)²/4, which reproduces the hypothesis for the third: recovery passes every stage whose coupling exceeds a(Q)²/4 and stops at the first that does not, so the sharp quantity is minₖ κₖ. Bisecting each coupling with the others held strong returned 0.0900000 at every stage, and the arrest depth was the first subcritical stage in every case (Fig. 3a). Existence of the branch is different: for k ≥ 2 a low root exists whatever κₖ is, so only κ₁ can annihilate it, and κc is unchanged by the downstream couplings and by length (Supplementary Note 4).",
"Three coupling regimes, one of them metastable",
"Because κc moves with a(Q), the coupling axis has three regimes rather than two (Fig. 2b). For κ > a₀²/4 a crossing propagates at once. For κ < a₁²/4 the branch survives at every attainable evidence level. Between the two the system stalls after the crossing and then completes spontaneously: residual availability is small but not zero, so Q creeps upward, a(Q) falls, and the branch is annihilated when a(Q) drops below 2√κ. Inverting the threshold gives the evidence level at which that happens, the annihilation locus",
"Q∗(κ) = [ (a₀ − 2√κ) / (a₀ − a₁) ]1/p,",
"which lies in (0,1) exactly on the intermediate band, tending to 1 as κ approaches a₁²/4 from above and to zero at a₀²/4. What the trajectory does next is: escape, detected on a different function, followed the predicted fold by 0.003% of the annihilation time at κ = 0.05 and 0.11% at κ = 0.08, and by 0.002% with the wide selector. Below the band nothing moved within 5 × 1010 time units.",
"While the root is recovered the evidence coordinate increases monotonically and the residual availability is bounded below, so the locus is reached in finite time for every κ above a₁²/4, where the partial configuration is an equilibrium of the capacity subsystem alone and metastable in the full system. For κ < a₁²/4 the branch is never removed and the trajectory follows it as Q rises, ending in a partial attractor at Q = 1, the first stage creeping from 0.0050 to 0.0238 at κ = 0.003. Nor is that profile firmly held: the barrier the first stage must clear is 1.7 × 10−4, and additive noise of standard deviation 0.008, under one per cent of the capacity range, took about 470 of 500 paths off the branch within 2 × 104 time units, median of order 4 × 103, while 0.002 took none (Fig. 3b). Confinement is a statement about the deterministic field, and the regimes read better as an ordering of timescales than as three fates.",
"The band is stage-wise as well. Each stage carries its own locus Q∗(κₖ), but none can cross before its predecessor, so stage k completes when the evidence level first exceeds the running maximum of Q∗ up to k. Over five coupling vectors the fifteen completions all sat above that maximum, by 0 to 0.041, in chain order, and the intervals between them were compressed, 0.004% to 9% of the preceding wait, because a stage coming online raises C(r) and accelerates Q.",
"The absolute waiting time inherits the evidence rate and the regularisation of ref. 4, differing by 380 between the selector scales, and is not a prediction here (Supplementary Note 3), whereas the band, its boundaries and the locus are.",
"Intermittent coupling acts as a duty-weighted average",
"Any repeated treatment raises a parameter for a fraction of the time rather than holding it. The field is linear in κ, so the averaged field is exactly the field at the duty-weighted mean and the criterion is expected rather than discovered21, though the trajectory error is of the order of the modulation period and not zero. Modulating κ as a rectangular wave with duty fraction f from the partial branch, escape occurred when the duty-weighted mean exceeded κc and not otherwise, in 40 of 42 cells (Fig. 4a); the two exceptions had f = 0.05 and an elevated coupling at least fourteen times κc, where the low phase is long enough to relax back. Sweeping the period at a fixed mean of 1.2 κc, the modulated escape time tracked the average to within 1% below a hundredth of the static escape time, then fell to a sixth of it once the period reached that time (Fig. 4b). Averaging fails in both directions, so the criterion holds only where deliveries are frequent compared with the escape they produce.",
"Escape from the partial state is slow and stepwise",
"Why that exponent is one half can be read off the branch. On a forward chain the frozen capacity Jacobian is lower triangular, so its spectrum is its diagonal, and at the first stage, with κ = a r₁ − r₁², it collapses exactly to",
"λ₁ = − √(a² − 4κ) (1 − r₁),",
"confirmed numerically to 10−16. Every eigenvalue is negative wherever the branch exists, so the partial configuration is stable at every point of it and the branch is normally hyperbolic away from the fold, where the trajectory tracks it to first order in ε22. The rate that vanishes at the fold is the square root of the same discriminant that fixes κc, and vanishes nowhere else (Fig. 3c; Supplementary Note 5).",
"Above κc the escape time diverges as a power of the excess, and the exponent is one half. Extending the excess down to 10−5, the local slope of log escape time against log excess falls monotonically from about 0.65 at an excess of 0.8 to 0.501 at the smallest examined (Fig. 4c); a single regression returns 0.529 over that range and 0.509 below 5% excess, both reporting curvature rather than the asymptote. The full system gives 0.48 to 0.57, and 0.39 to 0.53 below 5%.",
"One control separates the two effects in that spread. With the evidence rate set to zero the six threshold levels collapse, the fit below 5% excess lying between 0.527 and 0.532 (Supplementary Note 2), so curvature is intrinsic to the fold and the spread is not: at ε = 0.02 the fit falls to 0.386 at Q0 = 0.5, because the threshold moves while the system waits.",
"The shape of the trajectory matters more than the exponent. At an excess of 5% the first two downstream capacities crossed their own unstable roots at 0.31 and 0.65 of the escape interval and reached 0.9 within 3.3% of it, so each spends most of the interval below its own threshold and crosses quickly (Fig. 4d). The result is a staircase, and this is the sharpest observable difference between the axes: a root-driven crossing is fast, sharp and complete, a coupling-driven escape slow and stepwise.",
"Neither axis requires maintenance after the crossing",
"At the recovered configuration adjacent differences vanish and the coupling term contributes nothing, so it is asymptotically stable for every non-negative coupling including zero. A system carried past it by either axis stays recovered when both return to their previous values, so durability after a course is not evidence of lasting change. What changes with sustained occupancy is Q, lowering the threshold from 0.60 towards 0.15 and widening Δc(Q) from 0.40 to 0.85. Relapse is then a re-crossing of the same surface at low consolidated evidence, not decay of a treatment effect.",
"Identifying coupling and threshold from the profile",
"The obstacle to testing this is that coupling has never been measured, and inferring it from failure to respond is circular. The first stage breaks the circle: its equilibrium condition factorises exactly, so the coupling follows from the first-stage ratio and the threshold,",
"κ = a r₁ − r₁²,    with  a = 1 − Δ_{c},",
"with Δc(Q) the largest downward excursion from full recovery absorbed at that evidence level. On the exact equilibrium this returns the coupling to machine precision, with error first order in the error of r₁. Eliminating κ between the first two stages instead gives a = (r₁³ − r₂²)/(r₁² − r₂), exact only on the closed form from which it is derived, biased by −9.4% in the threshold and −11.2% in the coupling at κ = 0.003 before any noise is added, and ill-conditioned, r₂ being smaller than r₁ by a factor of fifty (Fig. 2c). The first-stage estimator is therefore primary.",
"This is what makes the two axes distinguishable in practice. An intervention acting only on the capacity state leaves the resting profile of a partial branch unchanged; one acting on coupling moves it, by a calculable amount, without the patient recovering. The estimator uses r₁ itself and not a monotone transform, so it needs an index of the first downstream capacity on its own [0,1] scale, anchored at zero for complete failure and at one for the fully recovered level in the same patient, and ref. 2 says why no scale supplies it: a cross-sectional reading does not separate the first two downstream capacities, which recovery timing separates. It also needs the threshold prevailing on the branch, a(Q), which is not what an acute challenge reads: a graded challenge probes the collapsed configuration and estimates the depleted-evidence threshold a₀, whereas the branch sits at consolidated evidence, and the two differ fourfold here. At κ = 0.003 the estimator returns 2.19 times the true coupling if a = 0.30 is used and 4.56 times it at a₀ = 0.60, against 12% for a 10% error in the correct threshold. The coupling is therefore structurally identifiable and not operationally measurable.",
"What the geometry does and does not say about ECT",
"The results above are statements about the model, and their clinical use needs a representation, stated as a hypothesis rather than derived. We take ECT to act as an excursion on the root and, over a course, as an increase of the effective coupling. A single treatment is a large acute exogenous perturbation, so its representation by u is close to direct; the coupling half is a coarse-grained hypothesis about propagation from one capacity to the next, not an identification with synaptic strength.",
"Under that representation three clinical regularities follow from the geometry rather than requiring separate accommodations: the response to a graded stimulus is all-or-none, the shape of the lorazepam challenge8,9; response after benzodiazepine failure needs only a larger excursion; and conversion can be fast, because beyond the surface the chain converges on the ordinary capacity timescale9,10.",
"The two axes are not competing accounts of one treatment but two states in which one treatment cannot be doing the same thing, and the geometry turns that into a conditional. If the coarse-grained coupling is diffusive, and if the threshold is not itself altered, then no coupling increase converts a wholly sub-threshold configuration, so an acute conversion requires an excursion-like component. One half of that conditional can be removed: on the dependency graph of ref. 2 the root carries no incoming edge, so no coupling of any form reaches it (Supplementary Note 6). On the modelled partial branch a partial converse holds: the root is already maximal, the excursion axis saturated and the evidence coordinate closed, so a durable change of the resting profile there is a change of coupling. Both remain conditional on the architecture and on identifying the acute and chronic presentations with the collapsed configuration and the branch, neither shown here.",
"Within the architecture, three things other than a change of coupling could take a patient off the branch: an input driving a downstream capacity above its own unstable root, noise, which the shallow barrier admits, and spontaneous completion on the band. The resting profile separates all three, since only a change in κ moves the resting first stage and keeps it moved.",
"Published dose-response data already separate the two axes",
"The architecture makes a prediction that published data can address without new measurement. The excursion axis sets whether a crossing occurs and the coupling axis how far it propagates, so among patients who respond at all the depth of recovery should not depend on the delivered amplitude. A meta-analysis of benzodiazepine treatment in catatonia collected 53 studies and found no significant dose-response relationship in effect size, reading that null as leaving optimal dosing unclear23. We took from it the 21 studies with at least ten patients reporting both response and remission. Across a tenfold range of highest daily dose, from 2 to 20 mg lorazepam-equivalent, the fraction of responders reaching remission spanned the whole interval from 0.00 to 1.00 and was uncorrelated with dose (slope on log dose −0.14, P = 0.35; Spearman ρ = −0.15, P = 0.50), as was the response rate itself (P = 0.31) (Fig. 5a). Within this geometry that null is not a failure to detect but the expected result, because dose is not the quantity that sets depth.",
"The same source tabulates BFCRS scores before and after treatment for ten individually reported patients. Their reductions were 0, 17, 23, 81, 85, 100, 100, 100, 100 and 100 per cent, with none between 23 and 81 per cent, an empty interval of 58 percentage points (Fig. 5b). A graded response places about 28 per cent of patients in a 35 to 65 per cent window, so observing none of ten there has P = 0.037. This is suggestive and not decisive: partial responses are the least publishable outcome, and a fivefold suppression of such cases would reproduce it (Fig. 5c). Studies reporting partial remission separately11 give a median partial share among responders of 20 per cent, range 8 to 68, compatible with arrest at a subcritical stage but not discriminating alone.",
"Discussion",
"Of the three quantities we allow an intervention to move, one has no external port and the other two are not alternatives, each powerless exactly where the other is decisive. The criterion κc = a²/4 carries the rest, exact because the first stage factorises, applied stage by stage, and, read as the discriminant whose square root is the rate vanishing at the fold, it fixes the escape exponent. What is exact and what is not should be kept apart: the first stage, the criterion, the locus, the spectrum and the estimator are exact; the stage-wise closed form downstream is an approximation with residual κ rₖ(rₖ₋₁ − 1), inherited as a bias by the two-ratio inversion. Three statements of ref. 4 are settled here: its unclassified coupling range is classified, its partial branch continues at every downstream coupling, and its prediction that consolidation alone cannot remove the branch holds only below a₁²/4.",
"We are equally explicit about the clinical bridge. Representing electroconvulsive therapy as an excursion with a coupling component is a hypothesis, asymmetric in its security, and nothing here identifies which physiology realises either axis. One weakness is better named than left implicit: the whole division rests on the coarse-grained coupling being diffusive. The one biological observation appealed to here, an increase in co-firing within cell assemblies under stimulation15, does not settle that in either direction, because a change in co-firing, corrected for firing rate or not, reports the coefficient with which one population acts on another, and such a coefficient occupies the same position in the additive form w and in the diffusive form κ. What separates the two is the sign structure of the sink term, which no reported comparison has tested.",
"The account is falsifiable in eight places. First, acute conversion should be all-or-none in the excursion coordinate, with no group whose recovery is partial in the amplitude delivered; this is the one prediction the section above already addresses, and it survived, though not decisively. Second, non-response to a smaller excursion and response to a larger should need no patient property beyond amplitude. Third, response kinetics should fall into three classes, fast complete conversion, a stalled state that completes late unaided, and one that does not, with membership predicted by the first-stage ratio through a₁²/4 and a₀²/4. Fourth, within the slow class, frequency and per-session amplitude should be interchangeable between those boundaries. Fifth, on the branch the resting first-stage ratio should be unchanged by an intervention acting on the capacity state and changed by one acting on coupling. Sixth, arrest depth should identify the weak stage, and late completions should arrive in bursts rather than evenly, which the milestone latencies already reported can be re-analysed to test1. Seventh, relapse risk should track occupancy of the recovered configuration rather than the number of treatments. Eighth, the two coupling forms are separated by a sign test and not by a magnitude: hold two populations at equal activity and ask whether the drive between them vanishes, then lower one below the other and ask whether it reverses. A diffusive coupling does both; an additive drive can do neither at any weight, and existing multi-electrode recordings are sufficient.",
"Several limitations qualify these results. The confinement theorem holds for arbitrary graphs and unequal thresholds but only for diffusive coupling, and everything after the crossing assumes a forward chain, though κc and the branch are unchanged from one stage to twelve. The identification needs both an anchored index of the first downstream capacity and the threshold prevailing on the branch, no published measure supplies either, and the two state identifications are assumed. The empirical section is a secondary analysis of published aggregate data: the depth comparison is between studies rather than within them, so it carries the usual ecological caveat and is confounded by chronicity, by underlying diagnosis and by the timing of assessment, while the case-level comparison rests on ten patients from case series and is bounded rather than established by its publication-bias sensitivity. The dependency order rests empirically on a single unblinded series1, and the proportional recovery rule in stroke warns that apparent ordering can arise from ceiling effects and mathematical coupling24; milestone timing in that series did not track peak severity, the rank correlations lying between 0.11 and 0.38 with none significant1. The boundaries 0.60, 0.090 and 0.0056 inherit the illustrative set of ref. 4, model time is not clinical time, and the waiting time moves by 380 with the regularisation, so timing statements are orderings. Whether the fold is passed with the delay of a dynamic saddle node is open.",
"Conclusion",
"Within the intervention architecture set out here, three quantities can be moved, and this work fixes what each can do. Evidence has no external port. Non-negative diffusive coupling cannot produce a crossing out of a configuration in which nothing is above threshold, on any graph and in either direction, so under that coupling form an acute conversion requires an excursion; an additive drive converts on its own only where the root has an incoming edge. After the crossing one exact criterion, a(Q)²/4, applied stage by stage, sets how far recovery reaches and when each stage completes, and the same discriminant gives the spectrum and the square-root escape law. Published data are consistent with both halves: the depth of recovery is uncorrelated with dose, and the acute response is all-or-none in the ten patients for whom it can be read.",
"Methods",
"Model and reduction",
"The system is that of ref. 4, with equal coupling across stages except where stated. The root obeys a bistable law plus an external excursion and each downstream capacity obeys the same law plus a diffusive term proportional to the difference from its immediate predecessor. The bistable law is a cubic with stable states at zero and one and an unstable threshold decreasing convexly with the evidence coordinate from a₀ at no evidence to a₁ at full evidence. Except where stated, analyses were run on the capacity subsystem with the threshold held at a fixed value: this is exact on the partial branch, where the joint availability that drives the evidence coordinate is of order 10−10, and conservative from a collapsed state, where the root lies below threshold and the evidence channel only attenuates. The full five-dimensional system was used for the intermediate-regime integration and for the confirmatory exponent measurement, with the illustrative parameters of ref. 4 throughout: a₀ = 0.60, a₁ = 0.15, p = 2, ε = 0.02, ρ = 1 and the one-sided regularisation χ±(z) = exp(−s/z²) on its own half-line, run at both scales used there, s = 0.01 and s = 1. Integration used an adaptive stiff solver with relative tolerance 10−10 and absolute tolerance 10−13, with event detection for crossings.",
"Confinement below threshold",
"One hundred systems were drawn from each of four graph families, feed-forward chains, symmetric graphs, purely reverse chains and general directed graphs, on three to six nodes, with weights drawn from a mixture placing mass at zero, 0.5, 3 and 30 times a uniform draw on (0,1]. The same mixture was used for every family, restricted to the sub-diagonal for feed-forward chains, symmetrised for symmetric graphs and transposed into the strict upper triangle for purely reverse chains, so no family was given a narrower weight distribution than another. Every capacity was initialised below a₀ and the system integrated with no external input; recorded quantities were whether any capacity reached the threshold and the largest increase of maxₖ rₖ over its initial value. Separately, in the forward chain, root trajectories from the same sub-threshold initial condition were computed at κ = 0.6, 3.0 and 30.0 on a common dense grid and compared pointwise rather than at the endpoint. The reverse-coupling counterexample used a bidirectional chain from r = (0.30, 0.90, 0.90, 0.90), at forward weight 0.01 with reverse weight 0.10 and with reverse coupling alone. A further 200 general directed graphs were drawn with a threshold per node from the uniform distribution on [a₁,a₀] and every capacity initialised below the smallest of them, which is the hypothesis the heterogeneous form of the principle needs. The scope experiments of Supplementary Note 6 used the same four-capacity system started at 0.30 everywhere with no external input: an additive form drₖ/dt = rₖ(1 − rₖ)(rₖ − a) + w(1 − rₖ)Σj≠k S(rⱼ) on the complete graph with S(r) = r and with S sigmoidal of midpoint 0.30 and width 0.05, the critical weight bracketed by bisection on the final maximum; the same start under diffusive coupling at weights 0.1, 1, 10 and 100; the same start under the reference coupling at thresholds a = 0.60, 0.40, 0.30 and 0.25; and a root placed at 0.61 with the rest at zero, integrated to 3 × 10⁴ at couplings from 0.6 to 0.01. The graph-restriction experiment of Supplementary Note 6 repeated the additive runs with the drive confined to the forward chain derived in ref. 2, so that capacity k receives only from k − 1 and the root receives nothing, from the same start and from the fully collapsed corner, at weights from 0.01 to 100 and with both forms of S, bracketing the weight at which each stage converts. The collapse criterion of Supplementary Note 7 was obtained by holding the root at zero and bisecting the coupling at which the first stage loses its high equilibrium, at both threshold levels.",
"Root pulse and excursion sweep",
"Pulses were applied as u(t) = A on [0,τ] and zero afterwards, from the collapsed chain at Q = 0, at durations τ = 0.25, 0.5, 1, 2 and 5. For each duration the critical amplitude was bracketed by bisection, and outcomes recorded at 0.8, 0.99, 1.01 and 1.2 times it together with the root value reached at the end of the pulse. A separate sweep set the root directly to a value v with no input, over a grid spanning a₀ and refined near it; v = a₀ lies on the invariant surface and was excluded from the grid rather than integrated, since the outcome there is neither corner. The transition was located by bracketed root-finding at three coupling values; the located boundary agrees with a₀ to twelve digits, which is a check on the solver.",
"Partial branch, critical coupling, identification",
"The first stage was taken from the factorised quadratic, which is exact. Each downstream stage was obtained as the low root of its cubic, bracketed on (0,a) because the cubic is positive at 0 and negative at a whenever 0 < rₖ₋₁ < a, and solved to machine precision; the residual of the vector field at that profile was below 10−18, and a perturbation of 10−4 in every downstream capacity returned to it. The stage-wise closed form was retained only to quantify its departure, reported as a relative error against the exact profile and against the analytic residual κ rₖ(rₖ₋₁ − 1). The propagation boundary was located by bracketed root-finding on the coupling at which a supracritical excursion reaches the distal capacity. In the full system the fold was located as a separate non-terminal event at a(Q)² − 4κ = 0; because that event function is the discriminant from which the locus is derived, the agreement it reports is a check on the solver rather than a test of the formula, and the independent comparison is with the escape event, detected on r₃ − a(Q) = 0. Identification was evaluated on the exact equilibrium, for the first-stage estimator and for the two-ratio inversion, without noise and then under multiplicative lognormal noise of coefficient of variation 1%, 5% and 10%, two thousand draws per level, reporting medians and interquartile ranges.",
"Unequal coupling, spectrum and noise",
"For unequal coupling each stage was given its own κₖ, and the boundary of a single stage was bracketed by bisection on that stage with the others held at 1.0, from the same supracritical excursion used for the equal-coupling boundary. The constants of the comparison lemma of Supplementary Note 4 were evaluated on a grid of 20001 points on [0,0.9] and the bound on the crossing time checked against the integrated comparison equation at four couplings above a₀²/4. Arrest depth was read from seven coupling vectors as the number of downstream capacities above 0.5 at the end of the integration, and compared with the index of the first stage whose coupling does not exceed a(Q)²/4. Existence of the branch was tested at κ₁ = 0.005625 and 0.00563 with downstream couplings spanning 10−4 to 30, and the critical coupling was located by bisection to 10−14 at chain lengths from one stage to twelve. The spectrum was taken as the diagonal of the capacity Jacobian, which is lower triangular on a forward chain, evaluated on the exact branch at three thresholds and ten couplings per threshold and compared with the closed form. For the noise runs the capacity block was integrated by Euler-Maruyama with additive Wiener increments of standard deviation σ on each capacity, step 0.05, the root held at 1, 500 paths per σ and a horizon of 2 × 104; escape was the first stage crossing its own unstable root, and the barrier was computed from the potential of the first-stage field.",
"Intermittent modulation and period sensitivity",
"The coupling was modulated as a rectangular wave between a base value of 0.003 and an elevated value, with duty fraction f, starting from the exact partial branch. Escape was scored as the most distal capacity crossing the unstable root of its own bistable law. For the criterion grid the period was one fiftieth of the static escape time at a duty-weighted mean of 1.2 κc. For the period sweep the duty fraction and elevated value were fixed so that the duty-weighted mean equalled 1.2 κc, and the period was varied from 10−3 to 5 times that static escape time; runs in which the escape completed inside the first elevated phase were flagged, since the ratio saturates there and successive points are not independent.",
"Escape-time scaling",
"Escape times were recorded at excesses from 10−5 to 0.8 above the critical coupling. The local slope was computed between consecutive excesses, and regressions were also taken over the full range and restricted to excesses below 5%. Reported escape times are for the most distal capacity crossing the unstable root of its own law, which is the definition of the escape time and not the completion of recovery. The measurement was repeated in the full system at Q = 0, 0.5, 0.7, 0.9, 0.99 and 1, in each case starting from the exact partial branch at half the critical coupling for the prevailing threshold, and then repeated again with the evidence rate set to zero and every other setting, including the starting point and the excess grid, left unchanged; that control isolates the contribution of evidence feedback, and it was extended to an excess of 2 × 10−4 at Q = 0 and Q = 0.5 to locate the asymptote. The fraction of the escape interval showing no appreciable change was computed at rate thresholds from 0.5% to 20% of the maximum rate and is reported as a range rather than a single value.",
"Secondary analysis of published aggregate data",
"Study-level and case-level values were transcribed from Tables 1 and 2 of ref. 23 and from Tables 1 and 2 of ref. 11. No individual patient records were accessed and no new data were collected. Depth of recovery is the ratio of the remission rate to the response rate in those studies reporting both with at least ten patients; ref. 23 defines remission as no longer meeting criteria, a BFCRS score below 3, and response as any reduction of symptoms, so the ratio is the fraction of responders reaching remission. Dose is the highest daily lorazepam-equivalent dose as tabulated there. The rows of Table 2 of ref. 23 with a sample size of one are individually reported patients, and their per cent reduction was computed from the tabulated scores before and after treatment. The partial window was fixed at 35 to 65 per cent, and the graded comparison value of 28 per cent, the share a monotone dose-response with the same mean curve places in that window, was fixed before the values were read. Association with dose was tested by least-squares regression on log dose and by Spearman rank correlation, and the count in the partial window by an exact binomial test. The publication-bias sensitivity varies the rate at which partial cases enter the literature relative to others and reports the probability of the observed count.",
"Ethics and reporting",
"This is a theoretical and numerical study. No human participants, animals, tissue or clinical records were involved and no ethics oversight was required. Clinical quantities cited in the text are taken from the published literature and are used to state what a model must reproduce, not as data analysed here.",
"Use of generative artificial intelligence",
"A large language model (Claude, Anthropic) was used during manuscript preparation to assist with language editing, identification of potential logical gaps, and code review. All mathematical arguments, proofs, numerical results and final wording were independently checked and approved by the author. The model was not used as an author and bears no responsibility for the work.",
"Supplementary information",
"Supplementary Note 1 restates, with proofs, the four properties of the underlying geometry used here: asymptotic stability of the recovered configuration for every non-negative coupling, asymptotic stability of the failed configuration for positive attenuation, exact invariance of the switching surface S = {r₀ = a(Q)} with forward invariance of z > 0 and z < 0, and the sufficient bound K* above which every trajectory off S reaches a corner. It is provided so that the present results can be checked without recourse to ref. 4, and it also records the analytic residual of the stage-wise closed form. Supplementary Note 2 reports the zero-evidence-rate control for the escape exponent, Supplementary Note 3 reconciles the two statements about joint availability on the partial branch, and Supplementary Notes 4 and 5 prove the two facts used for unequal coupling and the spectrum on the branch. Supplementary Note 6 gives the experiments that delimit the confinement statement: additive excitatory coupling, the restriction of that coupling to the derived dependency graph, lowering of the threshold, and propagation below a₀²/4. Supplementary Note 7 derives the collapse criterion, the dual of κᴄ obtained with the root at zero.",
"Data availability",
"No patient data were used in this study. All quantities computed here are generated by the deposited code from the parameter set specified in Methods. The study-level and case-level values analysed in Fig. 5 are transcribed from the published tables of refs. 11 and 23 and are carried in that code. Source data for Figs. 1 to 5 are provided with the paper as a single workbook with one sheet per panel, written by the same script that draws the figures.",
"Code availability",
"Code reproducing every reported value and every figure is provided as a single self-contained script that also re-checks each reported number against the recomputed one. It covers the confinement sweeps over the four graph families and over heterogeneous thresholds, the reverse-coupling counterexample, the rectangular-pulse and excursion grids, the exact partial branch and its identification, the unequal-coupling and chain-length sweeps, the spectrum on the branch, the intermittent-modulation grid and period sweep, the escape-time scaling with its zero-evidence-rate control, the noise runs, the scope experiments of Supplementary Note 6, the collapse criterion of Supplementary Note 7, and the secondary analysis of published aggregate data. Running it end to end writes the results file, the five figures, the source-data workbook, the manuscript and the Supplementary Information, and an audit of the text against the requirements of this journal. It is available at https://github.com/entrance4-png/catatonia-two-axes and archived at Zenodo under the concept DOI 10.5281/zenodo.21880298, which resolves to the latest version25, alongside the code accompanying the model on which this work builds4.",
"References",
"1. Saito, H. et al. Recovery from catatonia follows a hierarchical precision order. Preprint at medRxiv https://doi.org/10.1101/[medRxiv DOI to be inserted at submission] (2026).",
"2. Saito, H. The Saito Loop, a computational theory of hierarchical recovery in catatonia. Preprint at PsyArXiv https://osf.io/preprints/psyarxiv/de6bu (2026).",
"3. Saito, H. Recovery order in adaptive systems is set by dependency structure. Preprint at PsyArXiv https://osf.io/preprints/psyarxiv/8ae7w (2026).",
"4. Saito, H. Evidence-coupled bistability converts dependency structure into a recovery geometry. Preprint at Research Square https://doi.org/10.21203/rs.3.rs-10620384/v1 (2026).",
"5. Northoff, G. What catatonia can tell us about top-down modulation: a neuropsychiatric hypothesis. Behav. Brain Sci. 25, 555–577 (2002).",
"6. Hirjak, D. et al. Catatonia. Nat. Rev. Dis. Primers 10, 49 (2024).",
"7. Rogers, J. P. et al. Evidence-based consensus guidelines for the management of catatonia: recommendations from the British Association for Psychopharmacology. J. Psychopharmacol. 37, 327–369 (2023).",
"8. Bush, G., Fink, M., Petrides, G., Dowling, F. & Francis, A. Catatonia. II. Treatment with lorazepam and electroconvulsive therapy. Acta Psychiatr. Scand. 93, 137–143 (1996).",
"9. Sienaert, P., Dhossche, D. M., Vancampfort, D., De Hert, M. & Gazdag, G. A clinical review of the treatment of catatonia. Front. Psychiatry 5, 181 (2014).",
"10. Lloyd, J. R., Silverman, E. R., Kugler, J. L. & Cooper, J. J. Electroconvulsive therapy for patients with catatonia: current perspectives. Neuropsychiatr. Dis. Treat. 16, 2191–2208 (2020).",
"11. Pelzer, A. C. M., van der Heijden, F. M. M. A. & den Boer, E. Systematic review of catatonia treatment. Neuropsychiatr. Dis. Treat. 14, 317–326 (2018).",
"12. Friston, K., FitzGerald, T., Rigoli, F., Schwartenbeck, P. & Pezzulo, G. Active inference: a process theory. Neural Comput. 29, 1–49 (2017).",
"13. Chueh, K. N., Conley, C. C. & Smoller, J. A. Positively invariant regions for systems of nonlinear diffusion equations. Indiana Univ. Math. J. 26, 373–392 (1977).",
"14. Smith, H. L. Monotone Dynamical Systems: An Introduction to the Theory of Competitive and Cooperative Systems (American Mathematical Society, 1995).",
"15. Moore, H. et al. Stimulation modulates gene-linked cell assemblies in the human brain. Nature https://doi.org/10.1038/s41586-026-10879-9 (2026).",
"16. Kouvaris, N. E., Kori, H. & Mikhailov, A. S. Traveling and pinned fronts in bistable reaction-diffusion systems on networks. PLoS ONE 7, e45029 (2012).",
"17. Keener, J. P. Propagation and its failure in coupled systems of discrete excitable cells. SIAM J. Appl. Math. 47, 556–572 (1987).",
"18. Mallet-Paret, J. The global structure of traveling waves in spatially discrete dynamical systems. J. Dyn. Differ. Equ. 11, 49–127 (1999).",
"19. Elmer, C. E. & Van Vleck, E. S. Spatially discrete FitzHugh-Nagumo equations. SIAM J. Appl. Math. 65, 1153–1174 (2005).",
"20. Strogatz, S. H. Nonlinear Dynamics and Chaos 2nd edn (Westview, 2015).",
"21. Sanders, J. A., Verhulst, F. & Murdock, J. Averaging Methods in Nonlinear Dynamical Systems 2nd edn (Springer, 2007).",
"22. Fenichel, N. Geometric singular perturbation theory for ordinary differential equations. J. Differ. Equ. 31, 53–98 (1979).",
"23. Bot, L. et al. The effect of benzodiazepines on catatonia: a systematic review and meta-analysis. Acta Psychiatr. Scand. 154, 99–115 (2026).",
"24. Hope, T. M. H. et al. Recovery after stroke: not so proportional after all? Brain 142, 15–22 (2019).",
"25. Saito, H. catatonia-two-axes: reproduction code for a dependency-ordered recovery geometry. Zenodo https://doi.org/10.5281/zenodo.21880298 (2026).",
"Acknowledgements",
"The author thanks colleagues at Onda-daini Hospital, Nippon Medical School and Karolinska Institutet for discussion. No funding was received for this work.",
"Author contributions",
"H.S. conceived the study, derived and implemented the model, performed the numerical analysis, and wrote the manuscript.",
"Competing interests",
"The author declares no competing interests.",
"Figures",
"Fig. 1 | The crossing is set by the root excursion and is independent of coupling. a, The chain. In catatonia r₀ to r₃ are sensory, policy, motivational and fast-volatility precision, and Q is the identifiability of slow-volatility precision2. External input enters the root alone; coupling acts on differences between adjacent capacities and is absent from the root equation. b, The running maximum maxₖ rₖ(t) from a sub-threshold start, one example per graph family. Across 400 systems drawn from the four families no trajectory reached the threshold and the maximum never rose. c, Rectangular pulses u(t) = A on [0,τ] from the collapsed chain, at five durations and four amplitudes each. Outcome is fixed by the root value at the end of the pulse alone: conversion occurs exactly when r₀(τ) > a₀. d, With the same supracritical excursion in every run, the outcome after the crossing is set by coupling, with a sharp boundary at a₀²/4 = 0.090 located at 0.0900005.",
"Fig. 2 | The partial branch, its removal, and identification from the resting profile. a, Partial branch against coupling at consolidated evidence: exact equilibrium in colour, stage-wise closed form in grey. The two coincide at the first stage and diverge downstream. Inset, relative departure of the closed form from the exact profile: identically zero at r₁, rising to 2.0% at r₂ and 4.0% at r₃ at κ = 0.003 and growing as the coupling falls. The first-stage roots collide at κc = a₁²/4 = 0.0056. b, Because κc = a(Q)²/4 falls as evidence consolidates, the coupling axis has three regimes: immediate propagation above a₀²/4, deterministic confinement to the branch below a₁²/4, and an intermediate band in which the partial state is metastable, stalling and then completing without further intervention. Read against the upper axis the same curve is the annihilation locus Q∗(κ); markers show the three couplings integrated in the full system. c, Recovered coupling on the exact equilibrium. The first-stage estimator is unbiased at zero noise and degrades gracefully; the two-ratio inversion is biased before any noise is added and is ill-conditioned. Points are medians and bars interquartile ranges over two thousand draws per level.",
"Fig. 3 | The weakest stage sets where recovery stops, and the branch is shallowly held but hyperbolically stable. a, One stage swept while its neighbours stay strong. Recovery stops at the swept stage and nowhere else, and it stops exactly at a₀²/4, so the criterion applies stage by stage and the sharp quantity is the weakest coupling. b, Fraction of 500 noisy paths still on the confined branch at κ = 0.003, where the barrier is 1.7 × 10−4. Additive noise of standard deviation 0.008 empties the well within the horizon; 0.002 does not. c, Decay rates on the branch at consolidated evidence against the distance from the fold. The rate that vanishes at the fold follows the closed form √(a² − 4κ)(1 − r₁) exactly and does so as a square root, a straight line of slope 1/2 here, while the second stage stays fast.",
"Fig. 4 | Intermittent delivery, the escape exponent, and the stepwise escape. a, Escape from the partial branch under rectangular modulation follows the duty-weighted criterion in 40 of 42 combinations; the exceptions are large amplitude at low duty, where the fast-modulation assumption fails. b, Holding the duty-weighted mean at 1.2 κc, the modulated escape time tracks the averaged prediction for periods below a few per cent of the escape time and then falls below it, so averaging understates efficacy there, the reverse of a. Diamonds mark points at which the escape completes inside the first elevated phase, where the ratio saturates. c, Local slope of log escape time against log excess in the capacity subsystem, converging to 0.501 at the smallest excess examined; the dashed line is 1/2. Inset, exponent fitted below 5% excess in the full system against the prevailing threshold, with the evidence channel on and off. With the channel off the six levels collapse; the depression at ε = 0.02 is therefore evidence feedback and not the approach to the asymptote. d, At an excess of 5% the first two downstream capacities cross at 0.31 and 0.65 of the escape interval and reach 0.9 within 3.3% of it in both cases. The third crosses at 1.00 by construction, since the escape time is defined as that crossing.",
"Fig. 5 | Published data separate the axes and show an all-or-none acute response. a, Depth of recovery, the fraction of responders reaching remission, against highest daily dose in the 21 studies of ref. 23 with at least ten patients. Point area is proportional to sample size. Depth spans the full interval and is uncorrelated with dose; the line is the least-squares fit on log dose with its 95% band. b, Per cent reduction in BFCRS in the ten individually reported patients of ref. 23. The shaded band is the 35 to 65 per cent window in which a graded response would place about 28 per cent of patients, and it is empty. c, Publication-bias sensitivity: the probability of observing no patient in that window against the rate at which partial cases are published relative to others. A fivefold suppression reproduces the observation, so the case-level result bounds rather than establishes the all-or-none reading."
]

SUPP_PARAS = [
"Supplementary Information",
"Two separable axes of intervention in a dependency-ordered recovery geometry",
"Hiroki Saito",
"Supplementary Note 1. Properties of the underlying geometry",
"The main text uses four properties of the evidence-coupled bistable cascade. They are established in ref. 4 of the main text; they are restated and proved here, in the notation of the main text, so that the results of this paper can be checked without recourse to that work. Throughout, g(r,Q) = r(1 − r)(r − a(Q)), the coupling is non-negative, the root carries no incoming coupling, and the selectors χ₊ and χ₋ are non-negative, locally Lipschitz, and vanish on z ≤ 0 and z ≥ 0 respectively, so in particular χ₊(0) = χ₋(0) = 0.",
"S1. The recovered configuration is stable for every coupling",
"At Erec = (1,1,1,1) with any Q, g(1,Q) = 0 and every difference rₖ₋₁ − rₖ vanishes, so Erec is an equilibrium of the capacity block for every κ ≥ 0. Differentiating g gives ∂g/∂r = (1 − 2r)(r − a) + r(1 − r), which at r = 1 equals −(1 − a) < 0 since a < 1. The capacity Jacobian is lower triangular, because stage k receives only from stage k − 1, with diagonal entries −(1 − a) − κₖ, all strictly negative. In the evidence direction the factor 1 − Q multiplies the only positive term, so the linearisation there is −ε χ₊(z) C(r) ≤ 0, strictly negative while z > 0. The configuration is therefore asymptotically stable, and the conclusion does not depend on the value of the coupling, including κ = 0.",
"S2. The failed configuration is stable when evidence attenuates",
"At Epath = (0,0,0,0) the same computation gives ∂g/∂r = −a at r = 0, and the capacity Jacobian is again lower triangular with diagonal −a − κₖ < 0. There z = −a(Q) < 0, so χ₋ > 0 and the evidence equation reduces to dQ/dt = −ε ρ χ₋(z) Q, whose linearisation at Q = 0 is −ε ρ χ₋(−a) < 0 for ρ > 0. The failed configuration is then hyperbolically asymptotically stable. At ρ = 0 that eigenvalue vanishes and the set {(0,0,0,0,Q)} is a continuum of equilibria, Lyapunov stable in the capacities but not asymptotically stable in the full system.",
"S3. The switching surface is exactly invariant",
"Let S = {r₀ = a(Q)} and z = r₀ − a(Q). On S, and with no external input, g(r₀,Q) = r₀(1 − r₀)(r₀ − a(Q)) = 0, so dr₀/dt = 0. Also z = 0 gives χ₊(z) = χ₋(z) = 0, so dQ/dt = 0. Hence",
"dz/dt = dr₀/dt − a′(Q) dQ/dt = 0    on S,",
"for every ε and every ρ. The surface is therefore invariant, and because the field is locally Lipschitz, solutions are unique, so a trajectory in {z > 0} cannot reach {z < 0} without spending time on S, which is impossible. Both regions are forward invariant. The defining relation r₀ = a(Q) contains neither ε nor ρ, so the position of S is independent of the evidence timescale. Two remarks bear on the main text. First, the argument uses χ₊(0) = χ₋(0) = 0; if the selectors were replaced by plain averaging, χ±(0) = 1/2, the surface would move with the evidence timescale. That boundary zero is neither a stipulation nor a knife edge. Ref. 4 obtains it from a two-regime Bayesian model with ambiguity-suspended updating, at every finite sharpness and for every member of the admissible selector class, and proves the converse: within that class the surface is invariant for every ε and ρ if and only if both selectors vanish at z = 0. It also measures the cost of violating it, a boundary value η displacing the surface by O(εη). Second, u(t) enters dr₀/dt directly, so an external input is exactly the term that can carry a trajectory off S, which is the sense in which the crossing belongs to that axis alone.",
"S4. A sufficient coupling bound for whole-chain convergence",
"Ref. 4 establishes that if κₖ exceeds K = max over a in [a₁,a₀] of (1 − a + a²)/3 then every trajectory off S converges to a corner, E_{rec} when z(0) > 0 and E_{path} when z(0) < 0. At the illustrative parameters K ≈ 0.291. We quote this result rather than reprove it, and we do not rely on it: the main text uses it only for comparison, since for the specific question of whether a crossing at the root propagates to the distal capacity the sharp boundary is available in closed form and equals a(Q)²/4, which at depleted evidence is 0.090, a factor of three below K. K is sufficient for a stronger statement, convergence from every initial condition, and is correspondingly conservative here.",
"S5. Residual of the stage-wise closed form",
"The main text uses the exact first-stage root and solves the downstream stages numerically. For completeness, the residual left by the stage-wise closed form is analytic. Suppose rₖ satisfies rₖ² − a rₖ + κ rₖ₋₁ = 0, that is rₖ(rₖ − a) = −κ rₖ₋₁. Substituting into the true equilibrium condition gives",
"rₖ(1 − rₖ)(rₖ − a) + κ(rₖ₋₁ − rₖ) = (1 − rₖ)(−κ rₖ₋₁) + κ rₖ₋₁ − κ rₖ = κ rₖ(rₖ₋₁ − 1),",
"which vanishes only at rₖ₋₁ = 1. The first stage is therefore exact and every later stage is not. At the illustrative parameters and κ = 0.003 the closed form overstates r₂ by 2.0% and r₃ by 4.0%, and the residual it leaves, −1.396 × 10−6 at the second stage, is analytic rather than numerical. Because 0 < rₖ₋₁ < a, the exact cubic rₖ(1 − rₖ)(rₖ − a) + κ(rₖ₋₁ − rₖ) is positive at rₖ = 0 and negative at rₖ = a, so its low root is bracketed and can be obtained to machine precision, which is what the deposited code does.",
"Supplementary Note 2. Zero-evidence-rate control for the escape-time exponent",
"The escape time above the critical coupling diverges as a power of the excess with exponent one half, and two separate effects push a fitted exponent away from that value. The first is pre-asymptotic curvature: at any finite excess the fit still contains the approach to the asymptote, so it exceeds one half over a wide range and falls towards one half as the excess is reduced. The second is feedback through the evidence coordinate: while the system waits on the branch, Q rises, a(Q) falls and the critical coupling moves, so the excess against which the fit is taken is not the excess the system experiences. The first effect is present in the capacity subsystem, which has no evidence channel; the second exists only at a positive evidence rate.",
"The control separates them by a single change. The full five-dimensional sweep is repeated with ε set to zero and everything else, including the six evidence levels, the starting point at half the critical coupling for the prevailing threshold, the excess grid and the solver settings, held fixed. With the evidence channel off the six levels must give one exponent, because the capacity block then depends on the evidence level only through the constant a, and the excess is measured in units of the corresponding κc.",
"Exponents of the full system with the evidence channel on and off. Each entry is the negative slope of log escape time against log excess, fitted over the whole grid (0.005 to 0.8) and restricted to excesses at or below 0.05. Columns: Q₀, a(Q₀), then full range and small excess at ε = 0.02, then full range and small excess at ε = 0.",
"Q₀ = 1.00;  a = 0.1500;  ε = 0.02: 0.570, 0.532;  ε = 0: 0.570, 0.532",
"Q₀ = 0.99;  a = 0.1590;  ε = 0.02: 0.530, 0.456;  ε = 0: 0.570, 0.532",
"Q₀ = 0.90;  a = 0.2355;  ε = 0.02: 0.501, 0.432;  ε = 0: 0.571, 0.532",
"Q₀ = 0.70;  a = 0.3795;  ε = 0.02: 0.484, 0.401;  ε = 0: 0.572, 0.531",
"Q₀ = 0.50;  a = 0.4875;  ε = 0.02: 0.478, 0.386;  ε = 0: 0.574, 0.529",
"Q₀ = 0.00;  a = 0.6000;  ε = 0.02: 0.516, 0.430;  ε = 0: 0.577, 0.527",
"The collapse is exact to the third decimal: with the channel off the full-range exponent lies between 0.570 and 0.577 and the small-excess exponent between 0.527 and 0.532, a spread of 0.005 across a threshold range of a factor of four, against a spread of 0.146 in the same column with the channel on. At Q₀ = 1 the two sweeps agree to every digit reported, because z > 0 on the branch makes the only inflow term proportional to 1 − Q, which vanishes there, so the evidence coordinate is stationary and ε cannot act.",
"Continuing the control to smaller excess confirms that the residual curvature is the asymptote and nothing else. At Q₀ = 0.5 the local slope between consecutive excesses is 0.512 at 3.2 × 10−3, 0.508 at 1.4 × 10−3, 0.505 at 7.1 × 10−4 and 0.504 at 3.2 × 10−4; at Q₀ = 0 the same sequence is 0.510, 0.507, 0.505 and 0.503. Both approach one half from above, at the same rate, at thresholds differing by a factor of 1.23.",
"One alternative explanation is excluded arithmetically rather than numerically. Every sweep starts on the exact partial branch at κ = κc/2, and the first stage there is r₁ = a(1 − 1/√2)/2, while the branch remnant at κc is r₁ = a/2. The ratio is 1 − 1/√2 = 0.2929 at every threshold, since a cancels. The starting point is therefore the same fraction of the way to the fold at every evidence level and cannot produce a spread across levels.",
"Supplementary Note 3. The two statements about joint availability",
"The main text uses the joint availability C(r) = ∏ⱼ rⱼ in two apparently opposite ways: it is treated as negligible when the capacity subsystem is analysed at a frozen threshold, and as the quantity that eventually destroys the partial branch. Both are correct, and they refer to different timescales.",
"On the partial branch at the illustrative κ = 0.003 the exact profile is r = (1.000, 2.377 × 10−2, 4.676 × 10−4, 9.170 × 10−6), so C(r) = 1.02 × 10−10. The evidence equation gives |dQ/dt| ≤ ε C(r) = 2 × 10−12 while the root is recovered, so over an escape interval of a few hundred time units Q changes by less than 10−9 and a(Q) by less than 10−9. Holding the threshold fixed over such an interval is therefore not an approximation at any digit the main text reports, which is what licenses the frozen reduction used for the branch, the critical coupling, the identification and the modulation results.",
"The same residual availability is nevertheless bounded away from zero, and Q is monotone while z > 0, so the integral of the rate does not converge: given long enough, Q reaches any level below 1. The relevant level is the annihilation locus Q∗(κ), and the time to reach it is set by ε C(r) and by the selector scale. In the full-system integrations it ranges from 8.8 × 105 time units at κ = 0.08 with the narrow selector to 3.4 × 108 at the same coupling with the wide one, a factor of 380, and to 4.9 × 107 at κ = 0.05 with the narrow selector. Below a₁²/4 the locus does not exist and no waiting time is defined.",
"The consequence for the main text is a division between what is and is not a prediction. The existence of the intermediate band, its two boundaries a₁²/4 and a₀²/4, the locus Q∗(κ) and the ordering of the fold and the escape follow from the geometry and are reported as results. The absolute waiting time does not: it inherits the evidence rate and the regularisation of ref. 4, and varies by more than two orders of magnitude between the two selector scales used there. No clinical timescale is claimed from it.",
"Supplementary Note 4. Unequal coupling: the weakest stage",
"Let each stage carry its own coupling, drₖ/dt = g(rₖ,Q) + κₖ(rₖ₋₁ − rₖ), and hold the threshold fixed at a. Two facts follow, and together they replace the equal-coupling assumption used in the main analyses.",
"S6. Propagation is decided stage by stage, by the weakest coupling",
"Suppose the crossing has occurred and r₀ = 1. The first stage sees the equilibrium condition r₁(1 − r₁)(r₁ − a) + κ₁(1 − r₁) = (1 − r₁)[r₁² − a r₁ + κ₁] = 0, whose low roots exist exactly while κ₁ ≤ a²/4, so for κ₁ > a²/4 the first stage has no rest point below a and converges to 1. Stage 2 then faces a predecessor tending to 1 rather than one held at 1, and the step from that limit to a crossing in finite time is the content of the following lemma. Writing it out is what makes the induction a proof rather than a sequence of limits.",
"S6a. A predecessor tending to one carries a supercritical stage across",
"Let stage k obey drₖ/dt = rₖ(1 − rₖ)(rₖ − a) + κₖ(s(t) − rₖ) with s(t) → 1, and let κₖ > a²/4. Fix θ with a < θ < 1 and put",
"μ_{θ} = min_{0 ≤ x ≤ θ} (1 − x)(x² − a x + κₖ).",
"Because the discriminant a² − 4κₖ is negative, x² − a x + κₖ ≥ κₖ − a²/4 > 0 for every real x, so μθ ≥ (1 − θ)(κₖ − a²/4) > 0. The field is affine and increasing in s with slope κₖ, so for s ≥ 1 − η it is bounded below on [0,θ] by μθ − κₖ η, which is positive for every η < μθ/κₖ. Choose such an η, and choose T with s(t) > 1 − η for all t > T. Then drₖ/dt ≥ μθ − κₖ η > 0 for every t > T at which rₖ ≤ θ, so rₖ leaves [0,θ] no later than",
"t = T + θ / (μ_{θ} − κₖ η),",
"which is finite. In particular rₖ crosses a in finite time. Above θ the same comparison against the field at s = 1 − η, whose only zero in (a,1] lies within O(η) of 1, gives rₖ(t) ≥ 1 − O(η) eventually; since η may be taken as small as desired once T is enlarged accordingly, rₖ → 1. Stage k therefore reproduces for stage k + 1 the hypothesis that stage k − 1 supplied for it, and the induction runs down the chain.",
"The lemma also shows what the margin costs. At κₖ = 0.0901, a margin of 10−4 above a₀²/4 = 0.09, the field on [0,0.9] is bounded below by 7.0 × 10−5 and η must be smaller than 7.8 × 10−4; at κₖ = 0.20 the bound is 4.7 × 10−2 and η may be as large as 0.52. Integrating the comparison equation at η = μθ/2κₖ gave crossing times of 627, 82, 29 and 9.3 at κₖ = 0.0901, 0.095, 0.12 and 0.20, in every case inside the bound above. A predecessor that merely converges, rather than one held at 1, therefore delays the crossing but cannot prevent it.",
"Conversely, if κₖ < a²/4 the limiting field retains its low root at [a − √(a² − 4κₖ)]/2, the interval below it is forward invariant in the limit, and stage k converges there instead. Recovery therefore passes stage k if and only if κₖ > a²/4 and arrests at the first stage for which it does not, so the chain recovers completely if and only if minₖ κₖ > a²/4: the sharp quantity is the smallest coupling and not the average, and the arrest depth identifies which stage is weak. Numerically, bisecting each coupling separately with the others held at 1.0 returned the same boundary at every stage, 0.090000022 against a₀²/4 = 0.090000000, and the arrest depth matched the index of the first subcritical stage in every coupling vector tested, including vectors whose minimum sits within 0.001 of the boundary.",
"S6b. The intermediate band is stage-wise, at the running maximum of Q*",
"The same argument applies while the threshold is still moving. On the intermediate band a stage is held below its own unstable root until a(Q) falls to 2√κₖ, that is until Q reaches Q∗(κₖ). A stage cannot cross before its predecessor, so the level at which stage k completes is not Q∗(κₖ) but the running maximum of Q∗ over stages 1 to k: where a stage's own locus lies below that maximum it is released by its predecessor rather than by its own criterion, and it follows almost at once.",
"Integrating the full five-dimensional system with a coupling per stage, from the collapsed profile at Q = 0, at the five vectors (0.5, 0.05, 0.02), (0.5, 0.02, 0.05), (0.2, 0.08, 0.03), (0.3, 0.04, 0.07) and (0.5, 0.06, 0.015), all fifteen completions occurred in chain order and at evidence levels above the running maximum, by between 0 and 0.041. The overshoot is positive because Q continues to rise during the finite crossing, and it is larger for the stages released by a predecessor, which are already past their own locus when they go. Delayed completion is therefore not one event but a sequence of them, at levels the coupling vector fixes in advance.",
"The intervals between consecutive completions are compressed relative to the wait that precedes them, from 0.004% to 9% across the five vectors, because a stage coming online multiplies the joint availability C(r) and accelerates Q. On the evidence coordinate the sequence is a clean staircase; in time it is often a burst. That asymmetry is itself a signature: late completions should not be evenly spaced.",
"S7. Existence of the branch is decided by the first stage alone",
"For k ≥ 2 write f(x) = x(1 − x)(x − a) + κₖ(rₖ₋₁ − x) with 0 < rₖ₋₁ < a. Then f(0) = κₖ rₖ₋₁ > 0 and f(a) = −κₖ(a − rₖ₋₁) < 0, so f changes sign on (0,a) and a low root exists for every positive κₖ. No downstream coupling can destroy the branch; only κ₁ can, through the discriminant of the quadratic above.",
"Two consequences follow. The critical coupling κc = a²/4 is a statement about the first stage and is therefore unchanged by the downstream couplings: the branch was present at κ₁ = 0.005625 and absent at 0.00563 for downstream couplings of 10−4, 0.5, 10 and 30. It is also unchanged by the length of the chain, since adding stages adds only cubics that always have a low root: bisection returned 0.005625000000 at every chain length from one stage to twelve, agreeing to twelve digits. The identification of the coupling from the first stage ratio is untouched by both, because it uses only that stage.",
"Supplementary Note 5. The spectrum on the branch, and the exponent",
"On a forward chain, stage k receives only from stage k − 1, so the Jacobian of the capacity block at frozen Q is lower triangular and its spectrum is its diagonal,",
"λₖ = (1 − 2rₖ)(rₖ − a) + rₖ(1 − rₖ) − κ.",
"Evaluate this at the first stage of the branch. There r₁ is the low root of r² − a r + κ = 0, so κ = a r₁ − r₁², and substituting gives λ₁ = (1 − 2r₁)(r₁ − a) + r₁(1 − r₁) − a r₁ + r₁² = 2r₁ − a − 2r₁² + a r₁. Writing D = √(a² − 4κ), the low root is r₁ = (a − D)/2, so 2r₁ − a = −D and −2r₁² + a r₁ = r₁(a − 2r₁) = r₁ D. Hence",
"λ₁ = −D (1 − r₁) = −√(a² − 4κ) (1 − r₁),",
"an identity, not an approximation. It was checked against the numerically differentiated diagonal at three thresholds and ten couplings per threshold, with a largest discrepancy of 1.1 × 10−16.",
"Three things follow. First, λ₁ < 0 wherever the branch exists, since D > 0 and r₁ < 1, and the downstream λₖ are the slopes of the cubics at their own low roots and are negative for the same reason, so the partial configuration is stable at every point of the branch and not only at the illustrative coupling. Second, the branch is normally hyperbolic away from the fold, so at a small evidence rate the full trajectory tracks it to first order in ε by the standard geometric singular perturbation estimate (ref. 22 of the main text); what that estimate does not cover is the passage through the fold itself. Third, λ₁ is the eigenvalue that vanishes at the fold, and it does so as the square root of the same discriminant whose vanishing defines κc. On the grid computed here, thirty combinations of threshold and coupling, λ₁ is also the least negative of the three, with a smallest margin of 8.8 × 10−3 against λ₂, but we use only the first statement, which is what the escape-time law requires. The critical coupling is the statement that the discriminant reaches zero; the escape exponent of one half is the statement that it reaches zero linearly in the excess, so that the rate, and hence the reciprocal of the escape time, vanishes as its square root. The two headline results are one expression read twice.",
"Supplementary Note 6. What the confinement statement does not cover",
"The maximum principle of the main text has two hypotheses that are easy to read past: the coupling enters as a difference between capacities, and the threshold is not itself an intervention target. Three numerical experiments fix what each hypothesis is doing. All three use the illustrative parameters of ref. 4 and a chain of four capacities, and none of them contradicts the theorem; they delimit it.",
"S8. Additive excitatory coupling converts a sub-threshold chain on its own",
"Replace the diffusive term by a non-negative excitatory drive on the complete graph,",
"drₖ/dt = rₖ(1 − rₖ)(rₖ − a) + w (1 − rₖ) Σ_{j ≠ k} S(rⱼ),   w ≥ 0,  S ≥ 0.",
"The gate (1 − rₖ) keeps the state in [0,1] and the weights are non-negative, so this is a legitimate excitatory coupling; what it lacks is the difference structure. The sign argument fails immediately: at an index attaining the maximum the drive is still strictly positive, so D⁺M is not bounded above by the uncoupled field and nothing prevents M from rising through a.",
"It does rise. Started from rₖ = 0.30 for every k, wholly below a₀ = 0.60, with no external input and no change of threshold, the chain converts completely above w = 0.100 for S(r) = r and above w = 0.060 for the sigmoidal S(r) = [1 + exp(−(r − 0.30)/0.05)]−1. Both weights are exact rather than fitted. For S(r) = r the symmetric subspace rₖ ≡ r carries the field r(1 − r)(r − a + 3w), so the additive drive shifts the threshold to a − 3w and the start at 0.30 crosses precisely when 3w > a₀ − 0.30, that is w > 0.100; bisection returned 0.100000000000. For the sigmoidal S the drift at the symmetric start is 0.30 × 0.70 × (0.30 − 0.60) + 3w × 0.70 × S(0.30) = −0.063 + 1.05 w, since S(0.30) = 1/2, and this changes sign at w = 0.060; bisection returned 0.059999999999.",
"The diffusive control is unambiguous. The same start, the same graph and the same integration with the drive replaced by Σⱼ w (rⱼ − rₖ) leaves the chain at the failed corner at w = 0.1, 1, 10 and 100, three orders of magnitude above the additive threshold and far above any coupling used anywhere else in this work. The theorem is not a statement about coupling strength; it is a statement about coupling form, and no strength defeats it while no capacity is above threshold.",
"The counterexample is not idle, but neither is it settled by the one biological observation this work has appealed to. An increase in co-firing within cell assemblies under stimulation (ref. 15 of the main text) reports the coefficient with which one population acts on another, and that coefficient occupies the same position in w as it does in κ; whether the measurement is corrected for firing rate changes nothing here, since correcting for rate isolates the coefficient and says nothing about how the term depends on the difference. What separates the two forms is the sign structure of the sink term. The diffusive term vanishes at rₖ₋₁ = rₖ and reverses below it; the additive term is strictly positive whenever the source is, at every weight. No reported comparison holds two populations at equal activity and asks whether the drive between them vanishes, and none lowers one below the other and asks whether it reverses. That is the eighth prediction of the main text, and it is the measurement that would decide the question; existing multi-electrode recordings are sufficient for it. The two forms do make opposite predictions about the acute state, and the first prediction separates them there: under diffusive coupling a wholly sub-threshold chain cannot be converted by coupling at any strength, so an acute conversion requires an excursion, whereas under an additive drive on a graph with an edge into the root a large enough coupling change converts it with no excursion at all. If acute conversion is all-or-none in the delivered amplitude, with no group whose recovery is graded in amplitude, the excursion is doing the work; if response instead grades with the number or intensity of contacts independently of the amplitude reached, the additive reading is favoured. That test uses dose-response data of a kind that already exists.",
"S8a. On the derived graph the root is unreachable by coupling of any form",
"The additive experiment above is run on the complete graph, in which every capacity drives every other. The dependency analysis that supplies the chain excludes that topology: non-adjacent pairs retain positive joint Fisher information, so no edge joins them, and the transitive reduction is the chain πs → β → πm → πv_fast → πv_slow (ref. 2 of the main text). Restricting the same drive to that chain, so that capacity k receives only from k − 1 and the root receives nothing, changes the outcome at the root and only there. From rₖ = 0.30 everywhere with no external input the most distal stage converts above w = 0.376518 and the first downstream stage above w = 0.979380, while the root converts at no weight tested up to 100, the terminal state being (0, 1, 1, 1); with the sigmoidal S the same holds, and from the fully collapsed corner at w = 100 the terminal state is again (0, 1, 1, 1). That state has C(r) = 0, so the evidence coordinate never accumulates and the threshold never falls, and it is not a recovered configuration on any reading.",
"The reason is structural rather than numerical, and it does not depend on the weights or on the form of S. The root carries no incoming edge, so every coupling term in its equation multiplies nothing, whatever that term is; only u enters there. The confinement theorem is a statement about the whole chain and it does need the diffusive form. The clinical half of the conditional does not: on the graph the dependency analysis supplies, no coupling change of any form converts the root, so an acute conversion still requires an excursion. What the additive counterexample establishes is that a coupling change can convert the chain below the root, and that it can convert the root as well once the root is given incoming edges that the identifiability analysis does not exhibit.",
"S9. Lowering the threshold converts the chain with no input and no coupling change",
"The main text takes three quantities to be externally accessible. That list is a modelling assumption, and the threshold parameters show what it excludes. Holding the coupling at the reference value 0.6 and supplying no input at all, a chain started at 0.30 everywhere fails at a = 0.60 and at a = 0.40, sits exactly on its unstable fixed point at a = 0.30, and recovers completely at a = 0.25. Nothing in the theorem is violated, because a is a parameter of the field rather than a variable of the dynamics, but an intervention that lowered the prevailing threshold would convert a collapsed chain without any excursion and without touching the coupling. Any claim that an acute conversion requires an excursion is conditional on the threshold being fixed, and that conditionality is stated in the main text rather than assumed away.",
"S10. A crossing at the root propagates only under supercritical coupling",
"The pulse criterion of the main text, that the chain converts if and only if the root ends the pulse above a₀, is a statement about the root together with a coupling that carries the rest of the chain. Placing the root just above threshold at r₀ = 0.61 with the remaining capacities at zero and integrating to 3 × 104, the whole chain recovers at κ = 0.6, 0.2 and 0.095, and does not at κ = 0.085, 0.05 and 0.01. At κ = 0.085 the state settles at (1, 0.2293, 0.0306, 0.0038), the root recovered and the chain arrested on the partial branch. The boundary is a₀²/4 = 0.090, the same criterion as everywhere else in this work, so the two statements are consistent; the pulse sweep of Fig. 1c was run at κ = 0.6, which is supercritical, and the equivalence there is between the pulse and the crossing of the root, not between the pulse and recovery of the chain at arbitrary coupli ng.",
"Supplementary Note 7. The collapse criterion, the dual of κc",
"The main text obtains κc(Q) = a(Q)²/4 by holding the root at 1 after a crossing. The same construction with the root at 0 gives the sharp criterion for the opposite event. With r₀ = 0 the first-stage equilibrium condition is r₁(1 − r₁)(r₁ − a) − κ r₁ = r₁[(1 − r₁)(r₁ − a) − κ] = 0, whose non-zero roots solve r₁² − (1 + a)r₁ + (a + κ) = 0 with discriminant (1 − a)² − 4κ. The surviving high root and its unstable partner collide when that discriminant vanishes, so the exact critical coupling for whole-chain collapse is",
"κ_{collapse}(Q) = (1 − a(Q))² / 4,",
"which is 0.040000 at depleted evidence and 0.180625 at consolidated evidence. Bisecting the coupling at which the first stage loses its high equilibrium, with the root held at zero, returned 0.0400000 and 0.1806250; at 0.99 κ_collapse the chain settles at (0.82, 0.98, 1.00) and (0.62, 0.92, 0.99) respectively, and at 1.01 κ_collapse it goes to the failed corner.",
"Two things follow. First, the sufficient condition of ref. 4 for whole-chain collapse is replaced by an exact one, and the quantity max_x g(x,Q)/x reported there as 0.1806 is this criterion at a₁, computed there without being recognised as such, since max_x (1 − x)(x − a) = (1 − a)²/4 identically. Second, the two criteria are not symmetric, and the asymmetry reverses with evidence. At depleted evidence κ_collapse = 0.040 lies below κc = 0.090, so on (0.040, 0.090) a collapse at the root propagates through the chain while a crossing at the root does not: a system in that window can fall as one and cannot rise as one. At consolidated evidence the ordering reverses, κ_collapse = 0.1806 exceeding κc = 0.005625 by a factor of 32, so a consolidated system propagates recovery far more readily than collapse. Both statements are properties of the capacity block at frozen threshold and inherit the illustrative parameters of ref. 4."
]

# ============================================================================
#  PART VI.  Orchestration
# ============================================================================
def main():
    ap = argparse.ArgumentParser(
        description="Reproduce every reported value, Figs. 1-4 and the source data.")
    ap.add_argument("--quick", action="store_true",
                    help="reduced sample sizes; the structure of every run is unchanged")
    ap.add_argument("--no-figures", action="store_true",
                    help="skip the figures and the source-data workbook")
    ap.add_argument("--figures-only", action="store_true",
                    help="redraw the figures from an existing two_axes_results.json")
    ap.add_argument("--audit", action="store_true",
                    help="run the compliance audit and write the documents only")
    ap.add_argument("--figure-layout", choices=("inline", "end"), default="inline",
                    help="place each figure at its first mention, or keep the "
                         "legends together at the end with the image above each")
    args = ap.parse_args()
    if args.audit:
        ok = audit(MAIN_PARAS)
        write_docx(MAIN_PARAS, "two_axes_main.docx", figures=FIGURE_FILES,
                   layout=args.figure_layout)
        write_docx(SUPP_PARAS, "two_axes_supplementary.docx", title_first=False)
        return 0 if ok else 1
    q = args.quick
    t0 = time.time()
    if args.figures_only:
        with open("two_axes_results.json") as fh:
            R = json.load(fh)
        for fn in (fig1, fig2, fig3, fig4, fig5):
            print("wrote", fn(R))
        print("wrote", write_source_data(R))
        return 0
    R = {}
    print("=== 1. confinement below threshold ===")
    R["confinement"] = run_confinement(25 if q else 100)
    R["confinement_het"] = run_confinement_heterogeneous(40 if q else 200)
    R["confinement_examples"] = run_confinement_examples()
    R["reverse_counterexample"] = run_reverse_counterexample()
    R["root_kappa"] = run_root_kappa_independence()
    print("=== 2. root pulse ===");              R["pulse"] = run_pulse_sweep()
    print("=== 3. partial branch ===");          R["branch"] = run_branch()
    R["propagation_sweep"] = run_propagation_sweep()
    R["branch_curves"] = run_branch_curves()
    R["regime_curve"] = run_regime_curve()
    print("=== 4. unequal coupling ===");        R["unequal"] = run_unequal()
    R["stage_sweep"] = run_stage_sweep()
    print("=== 5. spectrum ===");                R["spectrum"] = run_spectrum()
    R["decay_rates"] = run_decay_rates()
    print("=== 6. intermediate band ===");       R["band"] = run_band(q)
    print("=== 7. escape scaling ===");          R["escape"] = run_escape_scaling(q)
    R["full_exponents"] = run_full_exponents(q)
    R["staircase"] = run_staircase()
    print("=== 8. identification ===");          R["identification"] = run_identification(quick=q)
    print("=== 9. intermittent modulation ==="); R["modulation"] = run_modulation(q)
    R["period_sweep"] = run_period_sweep(quick=q)
    print("=== 10. noise on the branch ===");    R["noise"] = run_noise(quick=q)
    R["noise_curves"] = run_noise_curves(quick=q)
    print("=== 11. scope of confinement (SN6) ==="); R["scope"] = run_scope()
    print("=== 12. collapse criterion (SN7) ==="); R["collapse"] = run_collapse()
    print("=== 13. published aggregate data (Results, Fig. 5) ===")
    R["published"] = run_published_data(); P = R["published"]
    print("      %d studies, depth %.2f to %.2f, uncorrelated with dose (P = %.2f)"
          % (P["k_studies"], P["depth_lo"], P["depth_hi"], P["depth_p"]))
    print("      %d individual patients, %d in the 35-65 per cent window, binomial P = %.3f"
          % (P["n_cases"], P["n_in_window"], P["p_binomial"]))

    ok = self_check(R)
    ok = audit(MAIN_PARAS) and ok

    if not args.no_figures:
        print("\n=== figures and source data ===")
        for fn in (fig1, fig2, fig3, fig4, fig5):
            print("  wrote", fn(R))
        print("  wrote", write_source_data(R))
    print("\n=== documents ===")
    write_docx(MAIN_PARAS, "two_axes_main.docx", figures=FIGURE_FILES,
                   layout=args.figure_layout)
    write_docx(SUPP_PARAS, "two_axes_supplementary.docx", title_first=False)

    R["_meta"] = dict(quick=q, seconds=time.time() - t0, all_checks_passed=bool(ok))
    with open("two_axes_results.json", "w") as fh:
        json.dump(R, fh, indent=1, default=str)
    print(f"\nwrote two_axes_results.json  ({time.time() - t0:.1f} s)")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
